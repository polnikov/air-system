import os
import platform
import math
import requests
import json
import webbrowser
import docx
from docx.shared import Cm, Pt, Mm
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from datetime import datetime
from functools import partial
from collections import deque
from numpy import array, around
from scipy.interpolate import RegularGridInterpolator, interp1d

from PySide6.QtCore import QSettings, QSize, Qt, QRegularExpression, QTimer, QStandardPaths
from PySide6.QtGui import QRegularExpressionValidator, QFont, QIcon, QRegularExpressionValidator, QAction
from PySide6.QtWidgets import (
    QApplication,
    QMainWindow,
    QMessageBox,
    QPushButton,
    QLabel,
    QLineEdit,
    QGridLayout,
    QVBoxLayout,
    QComboBox,
    QGroupBox,
    QRadioButton,
    QComboBox,
    QHBoxLayout,
    QWidget,
    QTabWidget,
    QScrollArea,
    QFileDialog,
    QProgressDialog,
    QMenu,
)

from constants import CONSTANTS


basedir = os.path.dirname(__file__)


try:
    from ctypes import windll  # Only exists on Windows.
    myappid = 'akudjatechnology.natural-air-system'
    windll.shell32.SetCurrentProcessExplicitAppUserModelID(myappid)
except ImportError:
    pass

version = '1.0.3'


class CustomComboBox(QComboBox):
    def showPopup(self):
        count = self.count()
        if count >= 30:
            self.view().setFixedHeight(self.view().sizeHintForRow(0) * 30)
        else:
            self.view().setFixedHeight(self.view().sizeHintForRow(0) * count)
        super().showPopup()


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.app_title = f'{CONSTANTS.APP_TITLE}_v{version}'
        self.setWindowTitle(self.app_title)
        self.box_style = 'QGroupBox::title { color: blue; }'
        self.table_style = '''
            QTableWidget {
                font-family: Consolas;
                font-size: 13px;
            }
            QHeaderView {
                font-family: Consolas;
                font-size: 13px;
            }
            QHeaderView::section {
                background-color: #F3F3F3
            }
        '''
        self.rows_count = 0
        self.current_file_path = None

        self.auto_save_timer = QTimer()
        self.auto_save_timer.timeout.connect(self.auto_save)
        self.auto_save_timer.start(300_000) # 5 minutes in milliseconds

        menubar = self.menuBar()
        file_menu = menubar.addMenu(CONSTANTS.MENU[0])

        manual_action = QAction(CONSTANTS.MENU[1], self)
        menubar.addAction(manual_action)
        manual_action.triggered.connect(self.open_manual)

        help_menu = menubar.addMenu(CONSTANTS.MENU[2])

        FILE_MENU_HANDLERS = (
            self.open,
            self.save,
            self.save_as,
            self.export,
        )
        for m in range(4):
            action = QAction(CONSTANTS.FILE_SUBMENU[m], self)
            action.setIcon(QIcon(os.path.join(basedir, CONSTANTS.FILE_MENU_ICONS[m])))
            action.setShortcut(CONSTANTS.FILE_MENU_SHORTCUTS[m])
            file_menu.addAction(action)
            action.triggered.connect(FILE_MENU_HANDLERS[m])
            if m in (0, 2):
                file_menu.addSeparator()

        HELP_MENU_HANDLERS = (
            self.show_about,
            self.check_updates,
        )
        for h in range(2):
            action = QAction(CONSTANTS.HELP_SUBMENU[h], self)
            action.setIcon(QIcon(os.path.join(basedir, CONSTANTS.HELP_MENU_ICONS[h])))
            help_menu.addAction(action)
            action.triggered.connect(HELP_MENU_HANDLERS[h])

        self.recent_files_menu = QMenu('Открыть недавние', self)
        self.recent_files_menu.setIcon(QIcon(os.path.join(basedir, CONSTANTS.RECENT_FILES_ICONS[-1])))
        file_menu.insertMenu(file_menu.actions()[1], self.recent_files_menu)
        self.recent_files = deque(maxlen=5)
        self.settings = QSettings('akudja.technology', 'natural-air-system')
        self.load_recent_files()
        self.update_recent_files_menu()

        menubar.setStyleSheet('''
            QMenuBar {
                font-family: Consolas;
                font-size: 12px;
                color: white;
                background: #606060 ;
            }
            QMenuBar::item:selected {
                color: black;
                background: #CCFFFF;
            }
            QMenu {
                font-family: Consolas;
                font-size: 12px;
                color: white;
                background: #606060 ;
            }
            QMenu::item:selected {
                color: black;
                background: #CCFFFF;
            }
            QMenu::separator {
                height: 1.5px;
            }
        ''')
        self.recent_files_menu.setStyleSheet('''
            QMenu {
                font-family: Consolas;
                font-size: 12px;
                color: white;
                background: #606060 ;
            }
            QMenu::item:selected {
                color: black;
                background: #CCFFFF;
            }
            QMenu::separator {
                height: 1.5px;
            }
        ''')

        self.tab_widget = QTabWidget(self)
        self.setCentralWidget(self.tab_widget)
        self.tab_widget.addTab(self.create_tab1_content(), CONSTANTS.TAB1_TITLE)
        self.tab_widget.addTab(self.create_tab2_content(), CONSTANTS.TAB2_TITLE)
        self.tab_widget.setTabVisible(1, False)

        self.showMaximized()
        self.setMaximumWidth(1680)


    def create_tab1_content(self) -> object:
        _widget = QWidget()
        _layout = QVBoxLayout()

        _hbox1 = QHBoxLayout()
        _hbox1.addWidget(self.create_init_data_box())
        _hbox1.addWidget(self.create_sputnik_calculation())

        _hbox2 = QHBoxLayout()
        _hbox2.setContentsMargins(-1, 3, -1, 0)
        _hbox2.addWidget(self.create_channel_cap())
        _hbox2.addWidget(self.create_buttons_box())
        _layout.setAlignment(_hbox2, Qt.AlignmentFlag.AlignTop)

        _hbox3 = QVBoxLayout()
        _hbox3.setContentsMargins(-1, 0, -1, 0)
        scroll_area = QScrollArea()
        self.scroll_area = scroll_area
        scroll_area.setStyleSheet('background-color: transparent; border: 0')
        scroll_area.setWidgetResizable(True)
        _hbox3.addWidget(scroll_area)
        scroll_widget = QWidget(scroll_area)
        scroll_area.setWidget(scroll_widget)
        self.main_box = QVBoxLayout(scroll_widget)
        self.main_box.setContentsMargins(1, 1, 1, 1)
        self.main_box.setSpacing(0)
        self.main_box.setAlignment(Qt.AlignmentFlag.AlignTop)
        self.main_box.addWidget(self.create_header())

        self.main_box.addWidget(self.create_last_row())
        for i in range(4):
            self.main_box.insertWidget(2, self.create_row())
        self.last_row.itemAtPosition(0, 0).widget().setText(self.get_sum_all_rows_str())
        self.change_dimensions_cells_in_table()

        _layout.addLayout(_hbox1)
        _layout.addLayout(_hbox2)
        _layout.addLayout(_hbox3)
        _widget.setLayout(_layout)
        return _widget


    def create_tab2_content(self) -> object:
        _widget = QWidget()
        _layout = QVBoxLayout()
        _hbox1 = QHBoxLayout()
        _hbox1.addWidget(self.create_deflector_calculation(), alignment=(Qt.AlignmentFlag.AlignTop | Qt.AlignmentFlag.AlignLeft))
        _layout.addLayout(_hbox1)
        _widget.setLayout(_layout)
        return _widget


    def create_init_data_box(self) -> object:
        _box = QGroupBox(CONSTANTS.INIT_DATA.TITLE)
        style = self.box_style
        _box.setStyleSheet(style)
        _box.setFixedWidth(450)

        self.init_data_layout = QGridLayout()
        _init_data = self.init_data_layout
        _init_data.setVerticalSpacing(10)
        labels = CONSTANTS.INIT_DATA.LABELS
        for i in range(len(labels)):
            label_0 = QLabel(labels[i][0])
            label_0.setFixedHeight(CONSTANTS.INIT_DATA.LINE_HEIGHT)
            label_0.setFixedWidth(320)
            line_edit = QLineEdit()
            line_edit.setAlignment(Qt.AlignmentFlag.AlignCenter)
            line_edit.setStyleSheet('QLineEdit { background-color: #E5FFCC; border: 1px solid #E2E2E2; border-radius: 5px; }')
            line_edit.setFixedWidth(CONSTANTS.INIT_DATA.INPUT_WIDTH)
            line_edit.setFixedHeight(CONSTANTS.INIT_DATA.LINE_HEIGHT)
            label_1 = QLabel(labels[i][1])
            label_1.setFixedHeight(CONSTANTS.INIT_DATA.LINE_HEIGHT)
            _init_data.addWidget(label_0, i, 0)
            _init_data.addWidget(line_edit, i, 1)
            _init_data.addWidget(label_1, i, 2)

        temperature_item = _init_data.itemAtPosition(0, 1)
        self.temperature_widget = temperature_item.widget()
        temperature_widget = self.temperature_widget
        temperature_widget.setObjectName('temperature')
        temperature_regex = QRegularExpression('^(?:\d|[12]\d|30)(?:\.\d)?$')
        temperature_validator = QRegularExpressionValidator(temperature_regex)
        temperature_widget.setValidator(temperature_validator)
        temperature_widget.textChanged.connect(self.calculate_sputnik_specific_pressure_loss)
        temperature_widget.textChanged.connect(self.calculate_sputnik_dynamic)
        temperature_widget.textChanged.connect(self.calculate_gravi_pressure)
        temperature_widget.textChanged.connect(self.calculate_dynamic)
        temperature_widget.textChanged.connect(self.calculate_specific_pressure_loss)
        temperature_widget.textChanged.connect(self.calculate_branch_pressure)
        temperature_widget.textChanged.connect(self.calculate_pass_pressure)
        temperature_widget.textChanged.connect(self.calculate_channel_cap)

        surface_item = _init_data.itemAtPosition(1, 1)
        self.surface_widget = surface_item.widget()
        surface_widget = self.surface_widget
        surface_widget.setObjectName('surface')
        surface_regex = QRegularExpression('^(?:[0-9]|[1-9]\d|100)(?:\.\d{1,3})?$')
        surface_validator = QRegularExpressionValidator(surface_regex)
        surface_widget.setValidator(surface_validator)
        surface_widget.setToolTip(CONSTANTS.INIT_DATA.SURFACE_INPUT_TOOLTIP)
        surface_widget.textChanged.connect(self.calculate_sputnik_specific_pressure_loss)
        surface_widget.textChanged.connect(self.calculate_specific_pressure_loss)

        floor_height_item = _init_data.itemAtPosition(2, 1)
        self.floor_height_widget = floor_height_item.widget()
        floor_height_widget = self.floor_height_widget
        floor_height_regex = QRegularExpression('^(?:[1-9]\d?|100)(?:\.\d{1,2})?$')
        floor_height_validator = QRegularExpressionValidator(floor_height_regex)
        floor_height_widget.setValidator(floor_height_validator)
        floor_height_widget.textChanged.connect(self.set_base_floor_height_in_table)
        floor_height_widget.textChanged.connect(self.calculate_height)

        self.channel_height_item = _init_data.itemAtPosition(3, 1)
        self.channel_height_widget = self.channel_height_item.widget()
        channel_height_widget = self.channel_height_widget
        channel_height_widget.setObjectName('channel_height')
        channel_height_regex = QRegularExpression('^(?:[1-9]|[1-9]\d|100)(?:\.\d{1,2})?$')
        channel_height_validator = QRegularExpressionValidator(channel_height_regex)
        channel_height_widget.setValidator(channel_height_validator)
        channel_height_widget.textChanged.connect(self.calculate_height)

        klapan_label = QLabel(CONSTANTS.INIT_DATA.KLAPAN_LABEL)
        self.klapan_widget = CustomComboBox()
        klapan_widget = self.klapan_widget
        klapan_widget.setObjectName('klapan_widget')
        klapan_widget.setStyleSheet('QComboBox { background-color: #E5FFCC; border: 1px solid #E2E2E2; border-radius: 5px; } QAbstractItemView { background-color: #E5FFCC }')
        klapan_widget.setFixedHeight(CONSTANTS.INIT_DATA.LINE_HEIGHT)
        klapan_widget.addItems(CONSTANTS.INIT_DATA.KLAPAN_ITEMS.keys())
        for i in (1, 3, 8, 14, 21, 25, 29, 38, 47, 51):
            klapan_widget.insertSeparator(i)
        klapan_widget.model().item(0).setEnabled(False)
        klapan_widget.setFixedWidth(170)
        self.klapan_widget_value = CONSTANTS.INIT_DATA.KLAPAN_ITEMS.get(klapan_widget.currentText())
        self.klapan_air_flow_label = QLabel(f'{self.klapan_widget_value} м<sup>3</sup>/ч')
        klapan_widget.currentTextChanged.connect(self.set_klapan_air_flow_in_label)
        klapan_widget.currentTextChanged.connect(self.calculate_sputnik_klapan_pressure_loss)
        klapan_widget.currentTextChanged.connect(self.activate_klapan_input)
        klapan_layout = QHBoxLayout()
        klapan_layout.addWidget(klapan_label)
        klapan_layout.addWidget(klapan_widget)

        klapan_input_label_1 = QLabel(CONSTANTS.INIT_DATA.KLAPAN_INPUT_LABEL_1)
        self.klapan_input = QLineEdit()
        klapan_input = self.klapan_input
        klapan_input.setObjectName('klapan_input')
        klapan_input_label_2 = QLabel('м<sup>3</sup>/ч')
        klapan_input.setFixedWidth(CONSTANTS.INIT_DATA.INPUT_WIDTH)
        klapan_input.setFixedHeight(CONSTANTS.INIT_DATA.LINE_HEIGHT)
        klapan_input.setAlignment(Qt.AlignmentFlag.AlignCenter)
        klapan_input.setStyleSheet('QLineEdit { background-color: #E0E0E0; border: 0; border-radius: 5px; }')
        klapan_input.setDisabled(True)
        klapan_input_regex = QRegularExpression('^(?:[1-9]|[1-9]\d|100)(?:)?$')
        klapan_input_validator = QRegularExpressionValidator(klapan_input_regex)
        klapan_input.setValidator(klapan_input_validator)
        klapan_input.setToolTip(CONSTANTS.INIT_DATA.KLAPAN_INPUT_TOOLTIP)
        klapan_input.textChanged.connect(self.calculate_sputnik_klapan_pressure_loss)

        _init_data.addLayout(klapan_layout, 4, 0, 1, 2)
        _init_data.addWidget(self.klapan_air_flow_label, 4, 2)
        _init_data.addWidget(klapan_input_label_1, 5, 0)
        _init_data.addWidget(self.klapan_input, 5, 1)
        _init_data.addWidget(klapan_input_label_2, 5, 2)
        _init_data.setColumnStretch(1, 1)
        _init_data.setColumnStretch(2, 1)
        _box.setLayout(_init_data)
        return _box


    def create_channel_cap(self) -> object:
        _widget = QWidget()
        self.channel_cap_widget = _widget
        _layout = QGridLayout()
        self.channel_cap_grid = _layout
        _layout.setObjectName(CONSTANTS.CAP.NAME)

        label_1 = QLabel(CONSTANTS.CAP.LABEL_1)
        label_1.setStyleSheet('QLabel { color: blue; }')
        label_1.setFixedHeight(CONSTANTS.CAP.LINE_HEIGHT)
        label_1.setFixedWidth(135)
        _layout.addWidget(label_1, 0, 0)

        self.cap_type = QComboBox()
        cap_type = self.cap_type
        cap_type.setStyleSheet('QComboBox { background-color: #E5FFCC; border: 1px solid #E2E2E2; border-radius: 5px; } QAbstractItemView { background-color: #E5FFCC }')
        cap_type.setFixedHeight(CONSTANTS.CAP.LINE_HEIGHT)
        cap_type.setFixedWidth(125)
        cap_type.addItems(CONSTANTS.CAP.TYPES)
        cap_type.model().item(0).setEnabled(False)

        cap_type.currentTextChanged.connect(self.change_channel_cap_visibility)
        cap_type.currentTextChanged.connect(self.set_channel_cap_tooltip)
        cap_type.currentTextChanged.connect(self.set_channel_cap_relations)
        cap_type.currentTextChanged.connect(self.calculate_channel_cap)
        cap_type.currentTextChanged.connect(self.calculate_full_pressure)
        cap_type.currentTextChanged.connect(self.calculate_full_pressure_last_row)

        cap_type.currentTextChanged.connect(self.show_deflector_in_table)
        cap_type.currentTextChanged.connect(self.calculate_available_pressure)
        cap_type.currentTextChanged.connect(self.activate_channel_cap)
        cap_type.currentTextChanged.connect(self.activate_deflector_tab)


        _layout.addWidget(cap_type, 0, 1)

        label_2 = QLabel('h')
        label_2.setFixedHeight(CONSTANTS.CAP.LINE_HEIGHT)
        label_2.setFixedWidth(10)
        label_2.setAlignment(Qt.AlignmentFlag.AlignCenter)
        label_2.hide()
        _layout.addWidget(label_2, 0, 2)

        self.input_h = QLineEdit()
        input = self.input_h
        input_regex = QRegularExpression(r'^0$|^1$|^2$|^[01]\.\d{1,2}$')
        input_validator = QRegularExpressionValidator(input_regex)
        input.setValidator(input_validator)
        input.setStyleSheet('QLineEdit { background-color: #E5FFCC; border: 1px solid #E2E2E2; border-radius: 5px; }')
        input.setFixedHeight(CONSTANTS.CAP.LINE_HEIGHT)
        input.setFixedWidth(40)
        input.setAlignment(Qt.AlignmentFlag.AlignCenter)
        input.setToolTip(CONSTANTS.CAP.INPUT_h_TOOLTIP)
        input.hide()

        input.textChanged.connect(self.calculate_channel_cap)
        _layout.addWidget(input, 0, 3)

        label_3 = QLabel('м')
        label_3.setFixedHeight(CONSTANTS.CAP.LINE_HEIGHT)
        label_3.setFixedWidth(10)
        label_3.hide()
        _layout.addWidget(label_3, 0, 4)

        label_4 = QLabel('h/Do')
        label_4.setFixedHeight(CONSTANTS.CAP.LINE_HEIGHT)
        label_4.setFixedWidth(30)
        label_4.setAlignment(Qt.AlignmentFlag.AlignCenter)
        label_4.hide()
        _layout.addWidget(label_4, 0, 5)

        self.fact_relation = QLabel()
        fact_relation = self.fact_relation
        fact_relation.setStyleSheet('QLabel { background-color: #EFEFEF; border: 0; border-radius: 5px; }')
        # pressure.setToolTip(CONSTANTS.BUTTONS.ADD_FLOOR_FOR_DELETE_TOOLTIP)
        fact_relation.setFixedHeight(CONSTANTS.CAP.LINE_HEIGHT)
        fact_relation.setFixedWidth(40)
        fact_relation.setAlignment(Qt.AlignmentFlag.AlignCenter)
        fact_relation.setToolTip(CONSTANTS.CAP.FACT_RELATION_TOOLTIP)
        fact_relation.hide()
        _layout.addWidget(fact_relation, 0, 6)

        self.relations = CustomComboBox()
        relations = self.relations
        relations.setStyleSheet('QComboBox { background-color: #E5FFCC; border: 1px solid #E2E2E2; border-radius: 5px; } QAbstractItemView { background-color: #E5FFCC }')
        relations.setFixedHeight(CONSTANTS.CAP.LINE_HEIGHT)
        relations.setFixedWidth(100)
        relations.hide()

        relations.currentTextChanged.connect(self.calculate_channel_cap)
        _layout.addWidget(relations, 0, 7)

        label_5 = QLabel('Pш')
        label_5.setFixedHeight(CONSTANTS.CAP.LINE_HEIGHT)
        label_5.setFixedWidth(15)
        label_5.setAlignment(Qt.AlignmentFlag.AlignCenter)
        label_5.hide()
        _layout.addWidget(label_5, 0, 8)

        self.cap_pressure = QLabel()
        pressure = self.cap_pressure
        pressure.setStyleSheet('QLabel { background-color: #EFEFEF; border: 0; border-radius: 5px; }')
        pressure.setFixedHeight(CONSTANTS.CAP.LINE_HEIGHT)
        pressure.setFixedWidth(50)
        pressure.setAlignment(Qt.AlignmentFlag.AlignCenter)
        pressure.setToolTip(CONSTANTS.CAP.PRESSURE_TOOLTIP)
        pressure.hide()
        _layout.addWidget(pressure, 0, 9)

        label_6 = QLabel('Па')
        label_6.setFixedHeight(CONSTANTS.CAP.LINE_HEIGHT)
        label_6.setFixedWidth(15)
        label_6.hide()
        _layout.addWidget(label_6, 0, 10)

        _widget.setMaximumWidth(650)
        _widget.setLayout(_layout)
        return _widget


    def create_buttons_box(self) -> object:
        _widget = QWidget()
        _widget.setFixedHeight(53)
        _layout = QHBoxLayout()
        _layout.addStretch()

        self.add_row_button = QPushButton()
        add_row_button = self.add_row_button
        add_row_button.setText(CONSTANTS.BUTTONS.ADD_BUTTON_TITLE)
        add_row_button.setFixedHeight(40)
        add_row_button.setFixedWidth(70)
        add_row_button.setStyleSheet('''
            QPushButton {
                background-color: #66CC00; border-radius: 5px;
            }
            QPushButton:hover {
                border: 2px solid grey;
                background: transparent;
            }
            QPushButton:pressed {
                border: 0px;
                background: black;
                color: white;
            }
        ''')
        _layout.addWidget(add_row_button)
        add_row_button.clicked.connect(self.add_row)
        add_row_button.clicked.connect(self.set_full_air_flow_in_deflector)
        add_row_button.clicked.connect(self.set_sputnik_airflow_in_table)
        add_row_button.clicked.connect(self.calculate_height)
        add_row_button.clicked.connect(self.change_dimensions_cells_in_table)
        add_row_button.clicked.connect(self.calculate_kms)
        add_row_button.clicked.connect(self.copy_table_dimensions)

        self.input_for_delete = QLineEdit()
        input = self.input_for_delete
        input_regex = QRegularExpression('^[1-9]{1,2}$|^100$')
        input_validator = QRegularExpressionValidator(input_regex)
        input.setValidator(input_validator)
        input.setStyleSheet('QLineEdit { background-color: #FFCCCC; border: 1px solid #E2E2E2; border-radius: 5px; }')
        input.setToolTip(CONSTANTS.BUTTONS.ADD_FLOOR_FOR_DELETE_TOOLTIP)
        input.setFixedHeight(40)
        input.setFixedWidth(40)
        input.setAlignment(Qt.AlignmentFlag.AlignCenter)
        _layout.addWidget(input)

        self.delete_row_button = QPushButton()
        delete_row_button = self.delete_row_button
        delete_row_button.setText(CONSTANTS.BUTTONS.DELETE_BUTTON_TITLE)
        delete_row_button.setFixedHeight(40)
        delete_row_button.setFixedWidth(70)
        delete_row_button.setStyleSheet('''
            QPushButton {
                background-color: #FF9999; border-radius: 5px;
            }
            QPushButton:hover {
                border: 2px solid grey;
                background: transparent;
            }
            QPushButton:pressed {
                border: 0px;
                background: black;
                color: white;
            }
        ''')
        _layout.addWidget(delete_row_button)
        delete_row_button.clicked.connect(self.delete_row)
        delete_row_button.clicked.connect(self.set_sputnik_airflow_in_table)
        delete_row_button.clicked.connect(self.set_full_air_flow_in_deflector)
        delete_row_button.clicked.connect(self.change_dimensions_cells_in_table)
        delete_row_button.clicked.connect(self.calculate_kms)

        _widget.setLayout(_layout)
        return _widget


    def create_sputnik_calculation(self) -> object:
        _box = QGroupBox(CONSTANTS.SPUTNIK_TABLE.TITLE)
        _box.setStyleSheet(self.box_style)
        _box.setFixedHeight(298)
        _box.setMinimumWidth(1200)
        _box.setMaximumWidth(1500)
        _layout = QGridLayout()
        _layout.setObjectName(CONSTANTS.SPUTNIK_TABLE.NAME)
        self.sputnik = _layout

        labels = CONSTANTS.SPUTNIK_TABLE.LABELS
        for i in range(len(labels)):
            label = QLabel(labels[i])
            label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            label.setStyleSheet('QLabel { background-color: #E0E0E0; border-radius: 5px; }')
            match i:
                case 0:
                    label.setMinimumWidth(CONSTANTS.SPUTNIK_TABLE.WIDTHS.get(i))
                case 1:
                    label.setMinimumWidth(CONSTANTS.SPUTNIK_TABLE.WIDTHS.get(i))
                case 3 | 4:
                    label.setMinimumWidth(CONSTANTS.SPUTNIK_TABLE.WIDTHS.get(i))
                case 14:
                    label.setMinimumWidth(CONSTANTS.SPUTNIK_TABLE.WIDTHS.get(i))
                case _:
                    label.setMinimumWidth(CONSTANTS.SPUTNIK_TABLE.WIDTHS.get('other'))
            label.setFixedHeight(CONSTANTS.SPUTNIK_TABLE.HEADER_HEIGHT)
            _layout.addWidget(label, 0, i)

        input_edit_style = 'QLineEdit { background-color: #E5FFCC; border: 1px solid #E2E2E2; border-radius: 5px; }'
        result_style = 'QLineEdit { background-color: #99CCFF; border: 0; border-radius: 5px; }'
        read_only_edit_style = 'QLineEdit { background-color: #EFEFEF; border: 0; border-radius: 5px; }'

        self.klapan_flow = QLineEdit()
        klapan_flow = self.klapan_flow
        klapan_flow.setAlignment(Qt.AlignmentFlag.AlignCenter)
        klapan_flow.setFixedHeight(CONSTANTS.SPUTNIK_TABLE.HEIGHT)
        klapan_flow.setMinimumWidth(CONSTANTS.SPUTNIK_TABLE.WIDTHS.get(1))
        klapan_flow.setToolTip(CONSTANTS.SPUTNIK_TABLE.KLAPAN_FLOW_TOOLTIP)
        klapan_flow.setStyleSheet(input_edit_style)
        klapan_flow.setObjectName('klapan_flow')
        _layout.addWidget(klapan_flow, 1, 1)

        klapan_flow.textChanged.connect(self.calculate_sputnik_klapan_pressure_loss)
        klapan_flow.textChanged.connect(self.calculate_full_pressure)

        for i in (1, 3, 5):
            edit = QLineEdit()
            edit.setAlignment(Qt.AlignmentFlag.AlignCenter)
            edit.setFixedHeight(CONSTANTS.SPUTNIK_TABLE.HEIGHT)
            edit.setMinimumWidth(CONSTANTS.SPUTNIK_TABLE.WIDTHS.get('other'))
            edit.setReadOnly(True)
            edit.setStyleSheet(result_style)
            _layout.addWidget(edit, i, 13)

        klapan = QLabel(CONSTANTS.SPUTNIK_TABLE.KLAPAN_LABEL)
        klapan.setAlignment(Qt.AlignmentFlag.AlignCenter)
        klapan.setFixedHeight(CONSTANTS.SPUTNIK_TABLE.HEIGHT)
        klapan.setMinimumWidth(CONSTANTS.SPUTNIK_TABLE.WIDTHS.get(0))
        klapan.setStyleSheet('QLabel { background-color: #E0E0E0; font-size: 11px; border-radius: 5px; }')
        _layout.addWidget(klapan, 1, 0)

        for i in (2, 4):
            label = QLabel(CONSTANTS.SPUTNIK_TABLE.SECTORS.get(i))
            label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            label.setStyleSheet('QLabel { background-color: #E0E0E0; border-radius: 5px;}')
            label.setMinimumWidth(CONSTANTS.SPUTNIK_TABLE.WIDTHS.get(0))
            label.setFixedHeight(CONSTANTS.SPUTNIK_TABLE.HEIGHT)
            _layout.addWidget(label, i, 0)

        for line in (2, 4):
            for i in range(1, 14):
                edit = QLineEdit()
                edit.setAlignment(Qt.AlignmentFlag.AlignCenter)
                edit.setMinimumWidth(CONSTANTS.SPUTNIK_TABLE.WIDTHS.get(i, 72))
                edit.setFixedHeight(CONSTANTS.SPUTNIK_TABLE.HEIGHT)
                if i in (1, 2, 3, 4, 11):
                    edit.setStyleSheet(input_edit_style)
                    match i:
                        case 1:
                            # Allows values: 0...200 with or without one digit after separator
                            regex = r'^(?:\d{1,2}|1\d{2}|2\d{2})(?:\.\d)?$'
                        case 2 | 11:
                            # Allows values: 0...100 with or without one | two digit after separator
                            regex = r'^(?:[0-9]|[1-9]\d|100)(?:\.\d{1,3})?$'
                        case 3 | 4:
                            # Allows values: whole numbers 0...2000
                            regex = r'^([1-9]\d{0,2}|1\d{3}|2000)?$'
                    validator = QRegularExpressionValidator(regex)
                    edit.setValidator(validator)
                else:
                    edit.setStyleSheet(read_only_edit_style)
                    edit.setReadOnly(True)

                match i:
                    case 2 | 7 | 8:
                        edit.textChanged.connect(self.calculate_sputnik_linear_pressure_loss)
                    case 3 | 4:
                        edit.textChanged.connect(self.calculate_sputnik_air_velocity)
                        edit.textChanged.connect(self.calculate_sputnik_diameter)
                        edit.textChanged.connect(self.calculate_sputnik_m)
                        edit.textChanged.connect(self.calculate_kms)
                        edit.textChanged.connect(self.copy_sputnik_dimensions)
                    case 5:
                        edit.textChanged.connect(self.calculate_sputnik_dynamic)
                        edit.textChanged.connect(self.calculate_sputnik_specific_pressure_loss)
                    case 6:
                        edit.textChanged.connect(self.calculate_sputnik_specific_pressure_loss)
                    case 10 | 11:
                        edit.textChanged.connect(self.calculate_sputnik_local_pressure_loss)
                        edit.textChanged.connect(self.calculate_sputnik_specific_pressure_loss)
                    case 9 | 12:
                        edit.textChanged.connect(self.calculate_sputnik_full_pressure_loss)
                _layout.addWidget(edit, line, i)
        for i in (3, 4):
            _layout.itemAtPosition(4, i).widget().setStyleSheet(read_only_edit_style)
            _layout.itemAtPosition(4, i).widget().setReadOnly(True)

        for i in (3, 5):
            radio_button = QRadioButton()
            radio_button.setToolTip(CONSTANTS.SPUTNIK_TABLE.RADIO_TOOLTIPS.get(i))
            _layout.addWidget(radio_button, i, 14, alignment=Qt.AlignmentFlag.AlignCenter)
        self.radio_button1 = _layout.itemAtPosition(3, 14).widget()
        self.radio_button2 = _layout.itemAtPosition(5, 14).widget()
        self.radio_button1.setChecked(True)
        
        self.radio_button1.clicked.connect(self.set_sputnik_airflow_in_table_by_radiobutton_1)
        self.radio_button2.clicked.connect(self.set_sputnik_airflow_in_table_by_radiobutton_2)

        self.radio_button1.clicked.connect(self.calculate_kms_by_radiobutton_1)
        self.radio_button2.clicked.connect(self.calculate_kms_by_radiobutton_2)
        self.radio_button1.clicked.connect(self.calculate_full_pressure_by_radiobutton_1)
        self.radio_button2.clicked.connect(self.calculate_full_pressure_by_radiobutton_2)
        self.radio_button1.clicked.connect(self.calculate_branch_pressure_by_radiobutton_1)
        self.radio_button2.clicked.connect(self.calculate_branch_pressure_by_radiobutton_2)

        self.radio_button2.clicked.connect(self.uncheck_radio_button_1)
        self.radio_button1.clicked.connect(self.uncheck_radio_button_2)

        cell_1_13 = _layout.itemAtPosition(1, 13).widget()
        cell_1_13.textChanged.connect(self.calculate_sputnik_result_pressure)

        cell_2_1 = _layout.itemAtPosition(2, 1).widget()
        cell_2_1.textChanged.connect(self.calculate_sputnik_air_velocity)
        cell_2_1.textChanged.connect(self.set_sputnik_airflow_in_table)
        cell_2_1.textChanged.connect(self.calculate_kms)

        cell_4_1 = _layout.itemAtPosition(4, 1).widget()
        cell_4_1.textChanged.connect(self.calculate_sputnik_air_velocity)
        cell_4_1.textChanged.connect(self.set_sputnik_airflow_in_table)
        cell_4_1.textChanged.connect(self.calculate_kms)
        cell_4_1.textChanged.connect(self.calculate_sputnik_result_pressure)

        cell_2_13 = _layout.itemAtPosition(2, 13).widget()
        cell_4_13 = _layout.itemAtPosition(4, 13).widget()
        cell_2_13.textChanged.connect(self.calculate_sputnik_result_pressure)
        cell_4_13.textChanged.connect(self.calculate_sputnik_result_pressure)

        self.cell_3_13 = _layout.itemAtPosition(3, 13).widget()
        self.cell_5_13 = _layout.itemAtPosition(5, 13).widget()
        self.cell_3_13.textChanged.connect(self.calculate_full_pressure)
        self.cell_5_13.textChanged.connect(self.calculate_full_pressure)

        cell_2_5 = _layout.itemAtPosition(2, 5).widget()
        cell_2_5.textChanged.connect(self.calculate_branch_pressure)
        cell_4_5 = _layout.itemAtPosition(4, 5).widget()
        cell_4_5.textChanged.connect(self.calculate_branch_pressure)

        _layout.setSpacing(3)
        _box.setLayout(_layout)
        return _box


    def copy_sputnik_dimensions(self, value) -> None:
        sputnik = self.sputnik
        a = sputnik.itemAtPosition(2, 3).widget().text()
        b = sputnik.itemAtPosition(2, 4).widget().text()
        sputnik.itemAtPosition(4, 3).widget().setText(a)
        sputnik.itemAtPosition(4, 4).widget().setText(b)


    def create_header(self) -> object:
        _widget = QWidget()
        _layout = QGridLayout()
        _layout.setObjectName(CONSTANTS.MAIN_TABLE.HEADER_NAME)
        self.header = _layout
        labels = CONSTANTS.MAIN_TABLE.LABELS

        for i in range(len(labels)):
            label = QLabel(labels[i])
            label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            match i:
                case 6:
                    label.setStyleSheet('QLabel { background-color: #CCCCFF; border-radius: 5px; }')
                case _:
                    label.setStyleSheet('QLabel { background-color: #E0E0E0; border-radius: 5px; }')

            match i:
                case 0:
                    label.setFixedWidth(CONSTANTS.MAIN_TABLE.WIDTHS.get(i))
                case 1 | 2 | 3 | 12 | 13 | 14 | 15 | 16 | 17 | 18 | 19 | 20:
                    label.setFixedWidth(CONSTANTS.MAIN_TABLE.WIDTHS.get(i))
                case 4 | 5 | 6 | 7 | 8 | 9:
                    label.setFixedWidth(CONSTANTS.MAIN_TABLE.WIDTHS.get(i))
                case 10 | 11:
                    label.setFixedWidth(CONSTANTS.MAIN_TABLE.WIDTHS.get(i))
            label.setFixedHeight(CONSTANTS.MAIN_TABLE.HEADER_HEIGHT)

            match i:
                case 4:
                    label.setToolTip(CONSTANTS.MAIN_TABLE.TOOLTIP_H)

            _layout.addWidget(label, 0, i)

        _layout.itemAtPosition(0, 6).widget().hide()

        _widget.setLayout(_layout)
        return _widget


    def create_row(self) -> object:
        self.rows_count += 1
        _widget = QWidget()
        _layout = QGridLayout()
        _layout.setContentsMargins(10, 3, 10, 3)
        _layout.setObjectName(CONSTANTS.MAIN_TABLE.ROW_NAME)

        input_edit_style = 'QLineEdit { background-color: #E5FFCC; border: 1px solid #E2E2E2; border-radius: 5px; }'
        read_only_edit_style = 'QLineEdit { background-color: #EFEFEF; border: 0; border-radius: 5px; }'

        for i in range(len(CONSTANTS.MAIN_TABLE.LABELS)):
            edit = QLineEdit()
            edit.setAlignment(Qt.AlignmentFlag.AlignCenter)
            edit.setFixedHeight(CONSTANTS.MAIN_TABLE.HEIGHT)
            match i:
                case 0:
                    edit.setFixedWidth(CONSTANTS.MAIN_TABLE.WIDTHS.get(i))
                case 1 | 2 | 3 | 12 | 13 | 14 | 15 | 16 | 17 | 18 | 19 | 20:
                    edit.setFixedWidth(CONSTANTS.MAIN_TABLE.WIDTHS.get(i))
                case 4 | 5 | 6 | 7 | 8 | 9:
                    edit.setFixedWidth(CONSTANTS.MAIN_TABLE.WIDTHS.get(i))
                case 10 | 11:
                    edit.setFixedWidth(CONSTANTS.MAIN_TABLE.WIDTHS.get(i))

            match i:
                case 0:
                    edit.setStyleSheet('QLineEdit { background-color: #E0E0E0; border: 0; border-radius: 5px; }')
                case 1 | 2 | 10 | 11:
                    edit.setStyleSheet(input_edit_style)
                case 6:
                    edit.setStyleSheet('QLineEdit { background-color: transparent; border: 0; border-radius: 5px; }')
                    edit.setReadOnly(True)
                case _:
                    edit.setReadOnly(True)
                    edit.setStyleSheet(read_only_edit_style)

            match i:
                case 1:
                    edit.textChanged.connect(self.calculate_height)
                    edit.textChanged.connect(self.calculate_linear_pressure_loss)
                case 3:
                    edit.textChanged.connect(self.set_full_air_flow_in_deflector)
                    edit.textChanged.connect(self.calculate_air_velocity)
                    edit.textChanged.connect(self.calculate_kms)
                    edit.textChanged.connect(self.calculate_channel_cap)
                case 4:
                    edit.textChanged.connect(self.calculate_gravi_pressure)
                case 5 | 6:
                    edit.textChanged.connect(self.calculate_available_pressure)
                case 7:
                    edit.textChanged.connect(self.calculate_result)
                case 20:
                    edit.textChanged.connect(self.calculate_result)
                case 8:
                    edit.textChanged.connect(self.calculate_pass_pressure)
                case 9:
                    edit.textChanged.connect(self.calculate_branch_pressure)
                case 10 | 11:
                    edit.textChanged.connect(self.calculate_air_velocity)
                    edit.textChanged.connect(self.calculate_diameter)
                    edit.textChanged.connect(self.calculate_m)
                    edit.textChanged.connect(self.calculate_kms)
                    edit.textChanged.connect(self.copy_table_dimensions)
                case 12:
                    edit.textChanged.connect(self.calculate_dynamic)
                    edit.textChanged.connect(self.calculate_specific_pressure_loss)
                    edit.textChanged.connect(self.calculate_pass_pressure)
                    edit.textChanged.connect(self.calculate_channel_cap)
                case 13:
                    edit.textChanged.connect(self.calculate_specific_pressure_loss)
                    edit.textChanged.connect(self.calculate_channel_cap)
                case 14 | 15:
                    edit.textChanged.connect(self.calculate_linear_pressure_loss)
                case 16 | 18 | 19:
                    edit.textChanged.connect(self.calculate_full_pressure)
                case 17:
                    edit.textChanged.connect(self.calculate_specific_pressure_loss)
                case 20:
                    edit.textChanged.connect(self.calculate_channel_cap)

            if i in (1, 10, 11):
                match i:
                    case 1:
                        # Allows values: 0...100 with or without one | two digit after separator
                        regex = r'^(?:[0-9]|[1-9]\d|100)(?:\.\d{1,3})?$'
                    case 10 | 11:
                        # Allows values: whole numbers 0...2000
                        regex = r'^([1-9]\d{0,2}|1\d{3}|2000)?$'
                validator = QRegularExpressionValidator(regex)
                edit.setValidator(validator)
            _layout.addWidget(edit, 0, i)

        _layout.itemAtPosition(0, 0).widget().setText(self.get_sum_all_rows_str())
        _layout.itemAtPosition(0, 6).widget().hide()

        _widget.setLayout(_layout)
        return _widget


    def create_last_row(self) -> object:
        _widget = QWidget()
        _layout = QGridLayout()
        _layout.setContentsMargins(10, 0, 10, 3)
        self.last_row = _layout
        _layout.setObjectName(CONSTANTS.MAIN_TABLE.LAST_ROW_NAME)

        input_edit_style = 'QLineEdit { background-color: #E5FFCC; border: 1px solid #E2E2E2; border-radius: 5px; }'
        read_only_edit_style = 'QLineEdit { background-color: #EFEFEF; border: 0; border-radius: 5px; }'

        for i in range(len(CONSTANTS.MAIN_TABLE.LABELS)):
            edit = QLineEdit()
            edit.setAlignment(Qt.AlignmentFlag.AlignCenter)
            edit.setFixedHeight(CONSTANTS.MAIN_TABLE.HEIGHT)
            match i:
                case 0:
                    edit.setFixedWidth(CONSTANTS.MAIN_TABLE.WIDTHS.get(i))
                case 1 | 2 | 3 | 12 | 13 | 14 | 15 | 16 | 17 | 18 | 19 | 20:
                    edit.setFixedWidth(CONSTANTS.MAIN_TABLE.WIDTHS.get(i))
                case 4 | 5 | 6 | 7 | 8 | 9:
                    edit.setFixedWidth(CONSTANTS.MAIN_TABLE.WIDTHS.get(i))
                case 10 | 11:
                    edit.setFixedWidth(CONSTANTS.MAIN_TABLE.WIDTHS.get(i))

            match i:
                case 0:
                    edit.setStyleSheet('QLineEdit { background-color: #FFCCCC; border: 0; border-radius: 5px; }')
                case 1 | 2 | 8 | 10 | 11:
                    edit.setStyleSheet(input_edit_style)
                case 6:
                    edit.setReadOnly(True)
                    edit.setStyleSheet('QLineEdit { background-color: #CCCCFF; border-radius: 5px; }')
                case _:
                    edit.setReadOnly(True)
                    edit.setStyleSheet(read_only_edit_style)

            match i:
                case 1:
                    edit.textChanged.connect(self.calculate_height)
                case 3:
                    edit.textChanged.connect(self.calculate_air_velocity)
                case 4:
                    edit.textChanged.connect(self.calculate_gravi_pressure)
                    edit.textChanged.connect(self.calculate_linear_pressure_loss_last_row)
                case 5 | 6:
                    edit.textChanged.connect(self.calculate_available_pressure)
                case 8:
                    edit.textChanged.connect(self.calculate_pass_pressure)
                case 7 | 20:
                    edit.textChanged.connect(self.calculate_result)
                case 9:
                    edit.textChanged.connect(self.calculate_branch_pressure)
                case 10 | 11:
                    edit.textChanged.connect(self.calculate_air_velocity)
                    edit.textChanged.connect(self.calculate_diameter)
                    edit.textChanged.connect(self.calculate_m)
                case 12:
                    edit.textChanged.connect(self.calculate_dynamic)
                    edit.textChanged.connect(self.calculate_specific_pressure_loss)
                    edit.textChanged.connect(self.calculate_pass_pressure)
                case 13:
                    edit.textChanged.connect(self.calculate_specific_pressure_loss)
                case 14 | 15:
                    edit.textChanged.connect(self.calculate_linear_pressure_loss_last_row)
                case 16 | 18 | 19:
                    edit.textChanged.connect(self.calculate_full_pressure_last_row)
                case 17:
                    edit.textChanged.connect(self.calculate_specific_pressure_loss)
                case 20:
                    edit.textChanged.connect(self.calculate_channel_cap)

            if i in (1, 8, 9, 10, 11):
                match i:
                    case 1:
                        # Allows values: 0...100 with or without one | two digit after separator
                        regex = r'^(?:[0-9]|[1-9]\d|100)(?:\.\d{1,3})?$'
                    case 8 | 9:
                        # Allows values: 0...100 with or without one | two digit after separator
                        regex = r'^(?:[0-9]|[1-9]\d|100)(?:\.\d{1,3})?$'
                    case 10 | 11:
                        # Allows values: whole numbers 0...2000
                        regex = r'^([1-9]\d{0,2}|1\d{3}|2000)?$'
                validator = QRegularExpressionValidator(regex)
                edit.setValidator(validator)
            _layout.addWidget(edit, 0, i)

        _layout.itemAtPosition(0, 0).widget().setText(self.get_sum_all_rows_str())
        _layout.itemAtPosition(0, 6).widget().hide()

        _widget.setLayout(_layout)
        return _widget


    def add_row(self) -> None:
        new_row = self.create_row()
        self.main_box.insertWidget(2, new_row)
        self.last_row.itemAtPosition(0, 0).widget().setText(self.get_sum_all_rows_str())
        self.set_base_floor_height_in_table()


    def delete_row(self) -> None:
        row_for_delete = self.input_for_delete.text()
        if all([row_for_delete, self.rows_count > 1]):
            row_for_delete = int(row_for_delete)
            if row_for_delete <= self.rows_count:
                rows = self.get_main_rows()
                row = rows.pop(len(rows) - row_for_delete)
                parent_widget = row.parent()
                parent_widget.setParent(None)
                parent_widget.deleteLater()
                self.rows_count -= 1
                self.input_for_delete.setText('')
                self.update_floor_number()
                self.calculate_height()
            elif row_for_delete == self.rows_count + 1:
                QMessageBox.critical(self, 'Ошибка', 'Последний этаж удалить нельзя')
            else:
                QMessageBox.critical(self, 'Ошибка', 'Такого этажа нет')
        elif all([row_for_delete, self.rows_count == 1]):
            QMessageBox.critical(self, 'Ошибка', 'Больше нет этажей для удаления')
        else:
            QMessageBox.critical(self, 'Ошибка', 'Не указан этаж для удаления')
            self.input_for_delete.setFocus()


    def uncheck_radio_button_1(self) -> None:
        self.radio_button2.setChecked(True)
        self.radio_button1.setChecked(False)


    def uncheck_radio_button_2(self) -> None:
        self.radio_button1.setChecked(True)
        self.radio_button2.setChecked(False)


    def calculate_sputnik_klapan_pressure_loss(self, value) -> None:
        name = self.sender().objectName()
        match name:
            case 'klapan_flow':
                current_klapan_flow = value
                hand_klapan_flow = self.klapan_input.text()
                klapan_widget_value = CONSTANTS.INIT_DATA.KLAPAN_ITEMS.get(self.klapan_widget.currentText())
                self._calculate_sputnik_klapan_pressure_loss(klapan_widget_value, hand_klapan_flow, current_klapan_flow)
            case 'klapan_widget':
                current_klapan_flow = self.klapan_flow.text()
                hand_klapan_flow = self.klapan_input.text()
                klapan_widget_value = CONSTANTS.INIT_DATA.KLAPAN_ITEMS.get(value)
                self._calculate_sputnik_klapan_pressure_loss(klapan_widget_value, hand_klapan_flow, current_klapan_flow)
            case 'klapan_input':
                current_klapan_flow = self.klapan_flow.text()
                hand_klapan_flow = value
                klapan_widget_value = CONSTANTS.INIT_DATA.KLAPAN_ITEMS.get(self.klapan_widget.currentText())
                self._calculate_sputnik_klapan_pressure_loss(klapan_widget_value, hand_klapan_flow, current_klapan_flow)


    def _calculate_sputnik_klapan_pressure_loss(self, klapan_widget_value, hand_klapan_flow, current_klapan_flow) -> None:
        item = self.sputnik.itemAtPosition(1, 13).widget()
        if hand_klapan_flow or klapan_widget_value == '--':
            klapan_flow = hand_klapan_flow
        else:
            klapan_flow = klapan_widget_value
        if all([current_klapan_flow, klapan_flow]):
            current_klapan_flow = int(current_klapan_flow)
            klapan_flow = int(klapan_flow)
            result = 10 * pow(current_klapan_flow / klapan_flow, 2)
            result = '{:.3f}'.format(round(result, 3))
            item.setText(result)
        else:
            item.setText('')


    def calculate_sputnik_air_velocity(self, value) -> None:
        sputnik = self.sputnik
        row, _, _, _ = sputnik.getItemPosition(sputnik.indexOf(self.sender()))

        air_flow = sputnik.itemAtPosition(row, 1).widget().text()
        a = sputnik.itemAtPosition(row, 3).widget().text()
        b = sputnik.itemAtPosition(row, 4).widget().text()
        if all([air_flow, a, b]):
            result = float(air_flow) / (3_600 * ((float(a) * float(b) / 1_000_000)))
            result = '{:.2f}'.format(round(result, 2))
            sputnik.itemAtPosition(row, 5).widget().setText(result)
        elif all([air_flow, a]):
            result = float(air_flow) / (3_600 * (3.1415 * (float(a) / 1_000) * (float(a) / 1_000) / 4))
            result = '{:.2f}'.format(round(result, 2))
            sputnik.itemAtPosition(row, 5).widget().setText(result)
        else:
            sputnik.itemAtPosition(row, 5).widget().setText('')


    def calculate_sputnik_diameter(self, value) -> None:
        sputnik = self.sputnik
        row, _, _, _ = sputnik.getItemPosition(sputnik.indexOf(self.sender()))

        a = sputnik.itemAtPosition(row, 3).widget().text()
        b = sputnik.itemAtPosition(row, 4).widget().text()
        if all([a, b]):
            diameter = (2 * float(a) * float(b) / (float(a) + float(b)) / 1_000)
            diameter = '{:.3f}'.format(round(diameter, 3))
            sputnik.itemAtPosition(row, 6).widget().setText(diameter)
        elif a:
            diameter = float(a) / 1_000
            diameter = '{:.3f}'.format(round(diameter, 3))
            sputnik.itemAtPosition(row, 6).widget().setText(diameter)
        else:
            sputnik.itemAtPosition(row, 6).widget().setText('')


    def calculate_sputnik_specific_pressure_loss(self, value) -> None:
        sputnik = self.sputnik
        temperature = self.temperature_widget.text()
        surface = self.surface_widget.text()
        for row in (2, 4):
            velocity = sputnik.itemAtPosition(row, 5).widget().text()
            diameter = sputnik.itemAtPosition(row, 6).widget().text()
            dynamic = sputnik.itemAtPosition(row, 10).widget().text()
            if all([temperature, diameter, velocity, surface, dynamic]):
                try:
                    temperature = float(temperature)
                    velocity = float(velocity)
                    diameter = float(diameter)
                    surface = float(surface)
                    dynamic = float(dynamic)
                    mu = 1.458 * pow(10, -6) * pow((273.15 + temperature), 1.5) / ((273.15 + temperature) + 110.4)
                    density = 353 / (273.15 + temperature)
                    v = mu / density
                    re = velocity * diameter / v
                    lam = 0.11 * pow(((surface / 1_000) / diameter + 68 / re), 0.25)
                    result = (lam / diameter) * dynamic
                    result = '{:.4f}'.format(round(result, 4))
                    sputnik.itemAtPosition(row, 7).widget().setText(result)
                except ZeroDivisionError:
                    sputnik.itemAtPosition(row, 7).widget().setText('')
            else:
                sputnik.itemAtPosition(row, 7).widget().setText('')


    def calculate_sputnik_dynamic(self, value) -> None:
        sputnik = self.sputnik
        row, _, _, _ = sputnik.getItemPosition(sputnik.indexOf(self.sender()))
        name = self.sender().objectName()

        if name == 'temperature':
            temperature = value
            for row in (2, 4):
                velocity = sputnik.itemAtPosition(row, 5).widget().text()
                if all([temperature, velocity]):
                    temperature = float(temperature)
                    velocity = float(velocity)
                    density = 353 / (273.15 + temperature)
                    result = velocity * velocity * density / 2
                    result = '{:.3f}'.format(round(result, 3))
                    sputnik.itemAtPosition(row, 10).widget().setText(result)
                else:
                    sputnik.itemAtPosition(row, 10).widget().setText('')
        else:
            temperature = self.temperature_widget.text()
            velocity = sputnik.itemAtPosition(row, 5).widget().text()
            if all([temperature, velocity]):
                temperature = float(temperature)
                velocity = float(velocity)
                density = 353 / (273.15 + temperature)
                result = velocity * velocity * density / 2
                result = '{:.3f}'.format(round(result, 3))
                sputnik.itemAtPosition(row, 10).widget().setText(result)
            else:
                sputnik.itemAtPosition(row, 10).widget().setText('')


    def calculate_sputnik_m(self, value) -> None:
        sputnik = self.sputnik
        row, _, _, _ = sputnik.getItemPosition(sputnik.indexOf(self.sender()))
        axis_x = CONSTANTS.REFERENCE_DATA.M.X
        axis_y = CONSTANTS.REFERENCE_DATA.M.Y
        z = array(CONSTANTS.REFERENCE_DATA.M.TABLE)
        interpolator = RegularGridInterpolator((axis_y, axis_x), z)

        a = sputnik.itemAtPosition(row, 3).widget().text()
        b = sputnik.itemAtPosition(row, 4).widget().text()
        if all([a, b]) and all([100 <= int(a) <= 1_500, 100 <= int(b) <= 1_500]):
            a, b = int(a), int(b)
            m = interpolator((b, a))
            m = '{:.3f}'.format(around(m, 3))
            sputnik.itemAtPosition(row, 8).widget().setText(m)
        elif a and 100 <= int(a) <= 1_500:
            a = int(a)
            m = interpolator((a, a))
            m = '{:.3f}'.format(around(m, 3))
            sputnik.itemAtPosition(row, 8).widget().setText(m)
        else:
            sputnik.itemAtPosition(row, 8).widget().setText('')


    def calculate_sputnik_linear_pressure_loss(self, value) -> None:
        sputnik = self.sputnik
        row, _, _, _ = sputnik.getItemPosition(sputnik.indexOf(self.sender()))

        l = sputnik.itemAtPosition(row, 2).widget().text()
        r = sputnik.itemAtPosition(row, 7).widget().text()
        m = sputnik.itemAtPosition(row, 8).widget().text()
        if all([l, r, m]):
            l = float(l)
            r = float(r)
            m = float(m)
            result = '{:.4f}'.format(round(l * r * m, 4))
            sputnik.itemAtPosition(row, 9).widget().setText(result)
        else:
            sputnik.itemAtPosition(row, 9).widget().setText('')


    def calculate_sputnik_local_pressure_loss(self, value) -> None:
        sputnik = self.sputnik
        row, _, _, _ = sputnik.getItemPosition(sputnik.indexOf(self.sender()))

        for row in (2, 4):
            dynamic = sputnik.itemAtPosition(row, 10).widget().text()
            kms = sputnik.itemAtPosition(row, 11).widget().text()
            if all([dynamic, kms]):
                dynamic = float(dynamic)
                kms = float(kms)
                result = '{:.3f}'.format(round(dynamic * kms, 3))
                sputnik.itemAtPosition(row, 12).widget().setText(result)
            else:
                sputnik.itemAtPosition(row, 12).widget().setText('')


    def calculate_sputnik_full_pressure_loss(self, value) -> None:
        sputnik = self.sputnik
        row, _, _, _ = sputnik.getItemPosition(sputnik.indexOf(self.sender()))

        for row in (2, 4):
            linear_pressure_loss = sputnik.itemAtPosition(row, 9).widget().text()
            local_pressure_loss = sputnik.itemAtPosition(row, 12).widget().text()
            if all([linear_pressure_loss, local_pressure_loss]):
                linear_pressure_loss = float(linear_pressure_loss)
                local_pressure_loss = float(local_pressure_loss)
                result = '{:.3f}'.format(round(linear_pressure_loss + local_pressure_loss, 3))
                sputnik.itemAtPosition(row, 13).widget().setText(result)
            else:
                sputnik.itemAtPosition(row, 13).widget().setText('')


    def calculate_sputnik_result_pressure(self, value) -> None:
        klapan_pressure = self.sputnik.itemAtPosition(1, 13).widget().text()
        one_side_pressure = self.sputnik.itemAtPosition(2, 13).widget().text()
        two_side_pressure = self.sputnik.itemAtPosition(4, 13).widget().text()
        if all([klapan_pressure, one_side_pressure]):
            klapan_pressure = float(klapan_pressure)
            one_side_pressure = float(one_side_pressure)
            first_result = '{:.3f}'.format(round((klapan_pressure + one_side_pressure), 3))
            self.sputnik.itemAtPosition(3, 13).widget().setText(first_result)
        else:
            self.sputnik.itemAtPosition(3, 13).widget().setText('')

        if all([klapan_pressure, two_side_pressure]):
            klapan_pressure = float(klapan_pressure)
            two_side_pressure = float(two_side_pressure)
            second_result = '{:.3f}'.format(round((klapan_pressure + two_side_pressure), 3))
            self.sputnik.itemAtPosition(5, 13).widget().setText(second_result)
        else:
            self.sputnik.itemAtPosition(5, 13).widget().setText('')


    def set_base_floor_height_in_table(self) -> None:
        base_floor_height = self.floor_height_widget.text()
        rows = self.get_all_rows()
        for row in rows:
            if base_floor_height:
                row.itemAtPosition(0, 1).widget().setText(base_floor_height)
            else:
                row.itemAtPosition(0, 1).widget().setText('')


    def change_dimensions_cells_in_table(self) -> None:
        rows = self.get_main_rows()
        for i in range(len(rows)):
            if i != len(rows) - 1:
                rows[i].itemAtPosition(0, 10).widget().setStyleSheet('QLineEdit { background-color: #EFEFEF; border: 0; border-radius: 5px; }')
                rows[i].itemAtPosition(0, 10).widget().setReadOnly(True)
                rows[i].itemAtPosition(0, 11).widget().setStyleSheet('QLineEdit { background-color: #EFEFEF; border: 0; border-radius: 5px; }')
                rows[i].itemAtPosition(0, 11).widget().setReadOnly(True)
            else:
                rows[i].itemAtPosition(0, 10).widget().setStyleSheet('QLineEdit { background-color: #E5FFCC; border: 1px solid #E2E2E2; border-radius: 5px; }')
                rows[i].itemAtPosition(0, 10).widget().setReadOnly(False)
                rows[i].itemAtPosition(0, 11).widget().setStyleSheet('QLineEdit { background-color: #E5FFCC; border: 1px solid #E2E2E2; border-radius: 5px; }')
                rows[i].itemAtPosition(0, 11).widget().setReadOnly(False)


    def copy_table_dimensions(self, value) -> None:
        rows = self.get_main_rows()
        a = rows[-1].itemAtPosition(0, 10).widget().text()
        b = rows[-1].itemAtPosition(0, 11).widget().text()
        for i in range(len(rows) - 1):
            if a:
                rows[i].itemAtPosition(0, 10).widget().setText(a)
            else:
                rows[i].itemAtPosition(0, 10).widget().setText('')
            if b:
                rows[i].itemAtPosition(0, 11).widget().setText(b)
            else:
                rows[i].itemAtPosition(0, 11).widget().setText('')


    def update_floor_number(self) -> None:
        rows = self.get_all_rows()
        for i in range(len(rows)):
            rows[i].itemAtPosition(0, 0).widget().setText(str(len(rows) - i))


    def set_sputnik_airflow_in_table_by_radiobutton_1(self, checked) -> None:
        if checked:
            one_side_flow = self.sputnik.itemAtPosition(2, 1).widget().text()
            if one_side_flow:
                self._fill_air_flow_column_in_main_table(one_side_flow)
            else:
                self._clean_air_flow_column_in_table()


    def set_sputnik_airflow_in_table_by_radiobutton_2(self, checked) -> None:
        if checked:
            one_side_flow = self.sputnik.itemAtPosition(2, 1).widget().text()
            two_side_flow = self.sputnik.itemAtPosition(4, 1).widget().text()
            if all([one_side_flow, two_side_flow]):
                flow = int(one_side_flow) + int(two_side_flow)
                max_flow = max(int(one_side_flow), int(two_side_flow))
                self._fill_air_flow_column_in_main_table(str(flow), str(max_flow))
            else:
                self._clean_air_flow_column_in_table()


    def set_sputnik_airflow_in_table(self, value) -> None:
        one_side_flow = self.sputnik.itemAtPosition(2, 1).widget().text()
        two_side_flow = self.sputnik.itemAtPosition(4, 1).widget().text()

        if all([self.radio_button1.isChecked(), one_side_flow]):
            self._fill_air_flow_column_in_main_table(one_side_flow)

        elif all([self.radio_button2.isChecked(), one_side_flow, two_side_flow]):
            flow = int(one_side_flow) + int(two_side_flow)
            max_flow = max(int(one_side_flow), int(two_side_flow))
            self._fill_air_flow_column_in_main_table(str(flow), str(max_flow))
        else:
            self._clean_air_flow_column_in_table()


    def _fill_air_flow_column_in_main_table(self, flow, max_flow=False) -> None:
        if max_flow:
            self.last_row.itemAtPosition(0, 3).widget().setText(max_flow)
        else:
            self.last_row.itemAtPosition(0, 3).widget().setText(flow)
        rows = self.get_main_rows()
        rows[-1].itemAtPosition(0, 3).widget().setText(flow)
        for i in range(len(rows)):
            result = int(flow) * int(rows[i].itemAtPosition(0, 0).widget().text())
            rows[i].itemAtPosition(0, 3).widget().setText(str(result))


    def _clean_air_flow_column_in_table(self) -> None:
        rows = self.get_all_rows()
        for row in rows:
            row.itemAtPosition(0, 3).widget().setText('')


    def calculate_height(self) -> None:
        channel_height = self.channel_height_widget.text()
        base_floor_height = self.floor_height_widget.text()
        if all([channel_height, base_floor_height]):
            rows = self.get_all_rows()
            result = '{:.2f}'.format(round(float(channel_height), 2))
            rows[-1].itemAtPosition(0, 4).widget().setText(result)

            for i in range(len(rows) - 2, -1, -1):
                floor_height = rows[i+1].itemAtPosition(0, 1).widget().text()
                previous_height = rows[i+1].itemAtPosition(0, 4).widget().text()

                if all([floor_height, previous_height]):
                    result = float(previous_height) - float(floor_height)
                    if (result > 0 and previous_height) or (result == 0 and previous_height):
                        result = '{:.2f}'.format(round(result, 2))
                        rows[i].itemAtPosition(0, 4).widget().setText(result)
                    else:
                        rows[i].itemAtPosition(0, 4).widget().setText('')
                else:
                    rows[i].itemAtPosition(0, 4).widget().setText('')
        else:
            rows = self.get_all_rows()
            for row in rows:
                row.itemAtPosition(0, 4).widget().setText('')


    def calculate_gravi_pressure(self) -> None:
        temperature = self.temperature_widget.text()
        rows = self.get_all_rows()
        for row in rows:
            height = row.itemAtPosition(0, 4).widget().text()
            if all([temperature, height]):
                temperature = float(temperature)
                height = float(height)
                result = CONSTANTS.ACCELERATION_OF_GRAVITY * height * ((353 / (273 + 5)) - (353 / (273 + temperature)))
                result = '{:.3f}'.format(round(result, 3))
                row.itemAtPosition(0, 5).widget().setText(result)
            else:
                row.itemAtPosition(0, 5).widget().setText('')


    def calculate_available_pressure(self) -> None:
        cap = self.cap_type.currentText()
        deflector_pressure = self.last_row.itemAtPosition(0, 6).widget().text()
        if all([cap == CONSTANTS.CAP.TYPES[-1], deflector_pressure]):
            rows = self.get_all_rows()
            for row in rows:
                gravi_pressure = row.itemAtPosition(0, 5).widget().text()
                if gravi_pressure:
                    gravi_pressure = float(row.itemAtPosition(0, 5).widget().text())
                    deflector_pressure = float(deflector_pressure)
                    result = 0.9 * gravi_pressure + deflector_pressure
                    result = '{:.3f}'.format(round(result, 3))
                    row.itemAtPosition(0, 7).widget().setText(result)
                else:
                    row.itemAtPosition(0, 7).widget().setText('')
        else:
            rows = self.get_all_rows()
            for row in rows:
                gravi_pressure = row.itemAtPosition(0, 5).widget().text()
                if gravi_pressure:
                    gravi_pressure = float(row.itemAtPosition(0, 5).widget().text())
                    result = 0.9 * gravi_pressure
                    result = '{:.3f}'.format(round(result, 3))
                    row.itemAtPosition(0, 7).widget().setText(result)
                else:
                    row.itemAtPosition(0, 7).widget().setText('')


    def set_deflector_pressure_in_table(self, value) -> None:
        if value:
            self.last_row.itemAtPosition(0, 6).widget().setText(value)
        else:
            self.last_row.itemAtPosition(0, 6).widget().setText('')


    def set_klapan_air_flow_in_label(self, value) -> None:
        klapan_flow = CONSTANTS.INIT_DATA.KLAPAN_ITEMS.get(value)
        self.klapan_air_flow_label.setText(f'{klapan_flow} м<sup>3</sup>/ч')


    def calculate_result(self, value) -> None:
        row = self.sender().parent().layout()
        available_pressure = row.itemAtPosition(0, 7).widget().text()
        full_pressure = row.itemAtPosition(0, 20).widget().text()

        if all([available_pressure, full_pressure]):
            available_pressure = float(available_pressure)
            full_pressure = float(full_pressure)
            if available_pressure > full_pressure:
                result = '[+] Тяга есть'
                row.itemAtPosition(0, 21).widget().setText(result)
                row.itemAtPosition(0, 21).widget().setAlignment(Qt.AlignmentFlag.AlignLeft)
                row.itemAtPosition(0, 21).widget().setStyleSheet('QLineEdit { background-color: #66CC00; border: 0; border-radius: 5px; }')
            else:
                result = '[-] Тяги нет'
                row.itemAtPosition(0, 21).widget().setText(result)
                row.itemAtPosition(0, 21).widget().setAlignment(Qt.AlignmentFlag.AlignLeft)
                row.itemAtPosition(0, 21).widget().setStyleSheet('QLineEdit { background-color: #FF3333; border: 0; border-radius: 5px; }')
        else:
            row.itemAtPosition(0, 21).widget().setText('')
            row.itemAtPosition(0, 21).widget().setStyleSheet('QLineEdit { background-color: #EFEFEF; border: 0; border-radius: 5px; }')


    def activate_klapan_input(self, value) -> None:
        if value == 'Другой':
            self.klapan_input.setDisabled(False)
            self.klapan_input.setStyleSheet('QLineEdit { background-color: #E5FFCC; border: 1px solid #E2E2E2; border-radius: 5px; }')
        else:
            self.klapan_input.setText('')
            self.klapan_input.setDisabled(True)
            self.klapan_input.setStyleSheet('QLineEdit { background-color: #E0E0E0; border: 0; border-radius: 5px; }')


    def create_deflector_calculation(self) -> object:
        _box = QGroupBox()
        _layout = QGridLayout()
        self.deflector = _layout

        labels = CONSTANTS.DEFLECTOR.LABELS
        for i in range(len(labels)):
            label = QLabel(labels[i])
            label.setFixedWidth(410)
            _layout.addWidget(label, i, 0)
            if i == 0:
                line_edit = QLineEdit()
                line_edit.setAlignment(Qt.AlignmentFlag.AlignCenter)
                line_edit.setFixedWidth(60)
                line_edit.setFixedHeight(30)
                line_edit.setStyleSheet('QLineEdit { background-color: #E5FFCC; border: 1px solid #E2E2E2; border-radius: 5px; }')
                _layout.addWidget(line_edit, 0, 1)
            else:
                line_label = QLineEdit()
                line_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
                line_label.setFixedWidth(60)
                line_label.setFixedHeight(30)
                line_label.setStyleSheet('QLineEdit { background-color: #EFEFEF; border: 0; border-radius: 5px; }')
                line_label.setReadOnly(True)
                _layout.addWidget(line_label, i, 1)

        wind_velocity = _layout.itemAtPosition(0, 1)
        self.wind_velocity = wind_velocity.widget()
        wind_velocity = self.wind_velocity
        wind_regex = QRegularExpression(r'^(?:[0-9]|[0-9]\d|50)(?:\.\d{1,2})?$')
        wind_validator = QRegularExpressionValidator(wind_regex)
        wind_velocity.setValidator(wind_validator)

        deflector_pressure = _layout.itemAtPosition(8, 1).widget()
        recommended_velocity = _layout.itemAtPosition(1, 1).widget()
        air_flow = _layout.itemAtPosition(2, 1).widget()
        required_deflector_square = _layout.itemAtPosition(3, 1).widget()
        deflector_diameter = _layout.itemAtPosition(4, 1).widget()
        real_deflector_velocity = _layout.itemAtPosition(5, 1).widget()
        velocity_relation = _layout.itemAtPosition(6, 1).widget()
        pressure_relation = _layout.itemAtPosition(7, 1).widget()
        deflector_pressure = _layout.itemAtPosition(8, 1).widget()
        deflector_pressure.setObjectName(CONSTANTS.DEFLECTOR.NAME)
        deflector_pressure.setStyleSheet('QLineEdit { background-color: #CCCCFF; border: 0; border-radius: 5px; }')

        wind_velocity.textChanged.connect(self.calculate_deflector_recommended_velocity)
        recommended_velocity.textChanged.connect(self.calculate_deflector_required_square)
        air_flow.textChanged.connect(self.calculate_deflector_required_square)
        required_deflector_square.textChanged.connect(self.calculate_deflector_diameter)
        deflector_diameter.textChanged.connect(self.calculate_deflector_real_velocity)
        air_flow.textChanged.connect(self.calculate_deflector_real_velocity)
        real_deflector_velocity.textChanged.connect(self.calculate_deflector_velocity_relation)
        wind_velocity.textChanged.connect(self.calculate_deflector_velocity_relation)
        velocity_relation.textChanged.connect(self.calculate_deflector_pressure_relation)
        wind_velocity.textChanged.connect(self.calculate_deflector_pressure)
        pressure_relation.textChanged.connect(self.calculate_deflector_pressure)

        deflector_pressure.textChanged.connect(self.set_deflector_pressure_in_table)

        _box.setLayout(_layout)
        return _box


    def set_full_air_flow_in_deflector(self, value) -> None:
        rows = self.get_all_rows()
        value = rows[0].itemAtPosition(0, 3).widget().text()

        if value:
            result = int(value) * len(rows)
            self.deflector.itemAtPosition(2, 1).widget().setText(str(result))
        else:
            self.deflector.itemAtPosition(2, 1).widget().setText('')


    def calculate_deflector_recommended_velocity(self, value) -> None:
        if value:
            wind_velocity = float(value)
            result = '{:.2f}'.format(round(wind_velocity * 0.3, 2))
            self.deflector.itemAtPosition(1, 1).widget().setText(result)
        else:
            self.deflector.itemAtPosition(1, 1).widget().setText('')


    def calculate_deflector_required_square(self, value) -> None:
        recommended_velocity = self.deflector.itemAtPosition(1, 1).widget().text()
        air_flow = self.deflector.itemAtPosition(2, 1).widget().text()
        if all([recommended_velocity, air_flow]):
            recommended_velocity = float(recommended_velocity)
            air_flow = float(air_flow)
            result = '{:.3f}'.format(round(air_flow / (3_600 * recommended_velocity), 3))
            self.deflector.itemAtPosition(3, 1).widget().setText(result)
        else:
            self.deflector.itemAtPosition(3, 1).widget().setText('')


    def calculate_deflector_diameter(self, value) -> None:
        required_deflector_square = value
        if required_deflector_square:
            required_deflector_square = float(required_deflector_square)
            for k in CONSTANTS.DEFLECTOR.DIAMETERS.keys():
                if required_deflector_square <= k:
                    diameter = CONSTANTS.DEFLECTOR.DIAMETERS.get(k)
                    self.deflector.itemAtPosition(4, 1).widget().setText(diameter)
                    break
                else:
                    self.deflector.itemAtPosition(4, 1).widget().setText('Нет значения!')
        else:
            self.deflector.itemAtPosition(4, 1).widget().setText('')


    def calculate_deflector_real_velocity(self, value) -> None:
        air_flow = self.deflector.itemAtPosition(2, 1).widget().text()
        diameter = self.deflector.itemAtPosition(4, 1).widget().text()
        if all([air_flow, diameter]) and diameter != 'Нет значения!':
            air_flow = float(air_flow)
            diameter = float(diameter)
            result = '{:.2f}'.format(round(air_flow / (3_600 * math.pi * (pow(diameter / 1_000, 2) / 4)), 2))
            self.deflector.itemAtPosition(5, 1).widget().setText(result)
        else:
            self.deflector.itemAtPosition(5, 1).widget().setText('')


    def calculate_deflector_velocity_relation(self, value) -> None:
        wind_velocity = self.deflector.itemAtPosition(0, 1).widget().text()
        real_velocity = self.deflector.itemAtPosition(5, 1).widget().text()
        if all([wind_velocity, real_velocity]):
            wind_velocity = float(wind_velocity)
            real_velocity = float(real_velocity)
            result = '{:.2f}'.format(round(real_velocity / wind_velocity, 2))
            self.deflector.itemAtPosition(6, 1).widget().setText(result)
        else:
            self.deflector.itemAtPosition(6, 1).widget().setText('')


    def calculate_deflector_pressure_relation(self, value) -> None:
        velocity_relation = value
        if velocity_relation:
            try:
                x = float(velocity_relation)
                f = interp1d(
                    CONSTANTS.REFERENCE_DATA.DEFLECTOR_PRESSURE_RELATION.X,
                    CONSTANTS.REFERENCE_DATA.DEFLECTOR_PRESSURE_RELATION.TABLE
                )
                result = '{:.2f}'.format(around(f(x), 2))
                self.deflector.itemAtPosition(7, 1).widget().setText(result)
            except ValueError:
                self.deflector.itemAtPosition(7, 1).widget().setText('')
        else:
            self.deflector.itemAtPosition(7, 1).widget().setText('')


    def calculate_deflector_pressure(self, value) -> None:
        wind_velocity = self.deflector.itemAtPosition(0, 1).widget().text()
        pressure_relation = self.deflector.itemAtPosition(7, 1).widget().text()
        if all([wind_velocity, pressure_relation]):
            wind_velocity = float(wind_velocity)
            pressure_relation = float(pressure_relation)
            result = pressure_relation * ((353 / (273.15 + 5)) * pow(wind_velocity, 2) / 2)
            result = '{:.3f}'.format(round(result, 3))
            self.deflector.itemAtPosition(8, 1).widget().setText(result)
        else:
            self.deflector.itemAtPosition(8, 1).widget().setText('')


    def activate_deflector_tab(self, text) -> None:
        if text == CONSTANTS.CAP.TYPES[-1]:
            self.tab_widget.setTabVisible(1, True)
        else:
            self.tab_widget.setTabVisible(1, False)


    def show_deflector_in_table(self, text) -> None:
        rows = self.get_all_rows()
        if text == CONSTANTS.CAP.TYPES[-1]:
            self.header.itemAtPosition(0, 6).widget().setVisible(True)
            for row in rows:
                row.itemAtPosition(0, 6).widget().setVisible(True)
        else:
            self.header.itemAtPosition(0, 6).widget().hide()
            for row in rows:
                row.itemAtPosition(0, 6).widget().hide()


    def calculate_air_velocity(self, value) -> None:
        row = self.sender().parent().layout()
        a = row.itemAtPosition(0, 10).widget().text()
        b = row.itemAtPosition(0, 11).widget().text()
        air_flow = row.itemAtPosition(0, 3).widget().text()

        if all([air_flow, a, b]):
            velocity = float(air_flow) / (3_600 * ((float(a) * float(b) / 1_000_000)))
            velocity = '{:.2f}'.format(round(velocity, 2))
            row.itemAtPosition(0, 12).widget().setText(velocity)
        elif all([air_flow, a]):
            velocity = float(air_flow) / (3_600 * (3.1415 * (float(a) / 1_000) * (float(a) / 1_000) / 4))
            velocity = '{:.2f}'.format(round(velocity, 2))
            row.itemAtPosition(0, 12).widget().setText(velocity)
        else:
            row.itemAtPosition(0, 12).widget().setText('')


    def calculate_diameter(self, value) -> None:
        row = self.sender().parent().layout()
        a = row.itemAtPosition(0, 10).widget().text()
        b = row.itemAtPosition(0, 11).widget().text()

        if all([a, b]):
            result = (2 * float(a) * float(b) / (float(a) + float(b)) / 1_000)
            result = '{:.3f}'.format(round(result, 3))
            row.itemAtPosition(0, 13).widget().setText(result)
        elif a:
            result = float(a) / 1_000
            result = '{:.3f}'.format(round(result, 3))
            row.itemAtPosition(0, 13).widget().setText(result)
        else:
            row.itemAtPosition(0, 13).widget().setText('')


    def calculate_dynamic(self, value) -> None:
        name = self.sender().objectName()
        if name == 'temperature':
            temperature = value
            rows = self.get_all_rows()
            for row in rows:
                velocity = row.itemAtPosition(0, 12).widget().text()
                if all([temperature, velocity]):
                    temperature = float(temperature)
                    velocity = float(velocity)
                    density = 353 / (273.15 + temperature)
                    result = velocity * velocity * density / 2
                    result = '{:.3f}'.format(round(result, 3))
                    row.itemAtPosition(0, 17).widget().setText(result)
                else:
                    row.itemAtPosition(0, 17).widget().setText('')
        else:
            temperature = self.temperature_widget.text()
            row = self.sender().parent().layout()
            velocity = row.itemAtPosition(0, 12).widget().text()
            if all([temperature, velocity]):
                temperature = float(temperature)
                velocity = float(velocity)
                density = 353 / (273.15 + temperature)
                result = velocity * velocity * density / 2
                result = '{:.3f}'.format(round(result, 3))
                row.itemAtPosition(0, 17).widget().setText(result)
            else:
                row.itemAtPosition(0, 17).widget().setText('')


    def calculate_specific_pressure_loss(self, value) -> None:
        name = self.sender().objectName()
        match name:
            case 'temperature' | 'surface':
                temperature = self.temperature_widget.text()
                surface = self.surface_widget.text()
                rows = self.get_all_rows()
                for row in rows:
                    velocity = row.itemAtPosition(0, 12).widget().text()
                    diameter = row.itemAtPosition(0, 13).widget().text()
                    dynamic = row.itemAtPosition(0, 17).widget().text()
                    if all([temperature, diameter, velocity, surface, dynamic]):
                        try:
                            temperature = float(temperature)
                            velocity = float(velocity)
                            diameter = float(diameter)
                            surface = float(surface)
                            dynamic = float(dynamic)
                            mu = 1.458 * pow(10, -6) * pow((273.15 + temperature), 1.5) / ((273.15 + temperature) + 110.4)
                            density = 353 / (273.15 + temperature)
                            v = mu / density
                            re = velocity * diameter / v
                            lam = 0.11 * pow(((surface / 1_000) / diameter + 68 / re), 0.25)
                            result = (lam / diameter) * dynamic
                            result = '{:.4f}'.format(round(result, 4))
                            row.itemAtPosition(0, 14).widget().setText(result)
                        except ZeroDivisionError:
                            row.itemAtPosition(0, 14).widget().setText('')
                    else:
                        row.itemAtPosition(0, 14).widget().setText('')
            case _:
                row = self.sender().parent().layout()
                temperature = self.temperature_widget.text()
                surface = self.surface_widget.text()
                velocity = row.itemAtPosition(0, 12).widget().text()
                diameter = row.itemAtPosition(0, 13).widget().text()
                dynamic = row.itemAtPosition(0, 17).widget().text()
                if all([temperature, diameter, velocity, surface, dynamic]):
                    try:
                        temperature = float(temperature)
                        velocity = float(velocity)
                        diameter = float(diameter)
                        surface = float(surface)
                        dynamic = float(dynamic)
                        mu = 1.458 * pow(10, -6) * pow((273.15 + temperature), 1.5) / ((273.15 + temperature) + 110.4)
                        density = 353 / (273.15 + temperature)
                        v = mu / density
                        re = velocity * diameter / v
                        lam = 0.11 * pow(((surface / 1_000) / diameter + 68 / re), 0.25)
                        result = (lam / diameter) * dynamic
                        result = '{:.4f}'.format(round(result, 4))
                        row.itemAtPosition(0, 14).widget().setText(result)
                    except ZeroDivisionError:
                        row.itemAtPosition(0, 14).widget().setText('')
                else:
                    row.itemAtPosition(0, 14).widget().setText('')


    def calculate_m(self, value) -> None:
        axis_x = CONSTANTS.REFERENCE_DATA.M.X
        axis_y = CONSTANTS.REFERENCE_DATA.M.Y
        z = array(CONSTANTS.REFERENCE_DATA.M.TABLE)
        interpolator = RegularGridInterpolator((axis_y, axis_x), z)

        row = self.sender().parent().layout()
        a = row.itemAtPosition(0, 10).widget().text()
        b = row.itemAtPosition(0, 11).widget().text()
        if all([a, b]) and all([100 <= int(a) <= 1_500, 100 <= int(b) <= 1_500]):
            a, b = int(a), int(b)
            m = interpolator((b, a))
            m = '{:.3f}'.format(around(m, 3))
            row.itemAtPosition(0, 15).widget().setText(m)
        elif a and 100 <= int(a) <= 1_500:
            a = int(a)
            m = interpolator((a, a))
            m = '{:.3f}'.format(around(m, 3))
            row.itemAtPosition(0, 15).widget().setText(m)
        else:
            row.itemAtPosition(0, 15).widget().setText('')


    def calculate_linear_pressure_loss(self, value) -> None:
        row = self.sender().parent().layout()
        l = row.itemAtPosition(0, 1).widget().text()
        r = row.itemAtPosition(0, 14).widget().text()
        m = row.itemAtPosition(0, 15).widget().text()
        if all([l, r, m]):
            l = float(l)
            r = float(r)
            m = float(m)
            result = '{:.4f}'.format(round(l * r * m, 4))
            row.itemAtPosition(0, 16).widget().setText(result)
        else:
            row.itemAtPosition(0, 16).widget().setText('')


    def calculate_linear_pressure_loss_last_row(self, value) -> None:
        row = self.sender().parent().layout()
        l = row.itemAtPosition(0, 4).widget().text()
        r = row.itemAtPosition(0, 14).widget().text()
        m = row.itemAtPosition(0, 15).widget().text()
        if all([l, r, m]):
            l = float(l)
            r = float(r)
            m = float(m)
            result = '{:.4f}'.format(round(l * r * m, 4))
            row.itemAtPosition(0, 16).widget().setText(result)
        else:
            row.itemAtPosition(0, 16).widget().setText('')


    def calculate_kms_by_radiobutton_1(self, checked) -> None:
        if checked:
            one_side_flow = self.sputnik.itemAtPosition(2, 1).widget().text()
            sputnik_a = self.sputnik.itemAtPosition(2, 3).widget().text()
            sputnik_b = self.sputnik.itemAtPosition(2, 4).widget().text()
            if sputnik_b == '':
                sputnik_b = sputnik_a
            if all([one_side_flow, sputnik_a, sputnik_b]):
                self._calculate_kms_by_radiobutton(one_side_flow, sputnik_a, sputnik_b)


    def calculate_kms_by_radiobutton_2(self, checked) -> None:
        if checked:
            one_side_flow = self.sputnik.itemAtPosition(2, 1).widget().text()
            two_side_flow = self.sputnik.itemAtPosition(4, 1).widget().text()
            if all([one_side_flow, two_side_flow]):
                flow = str(max(float(one_side_flow),  float(two_side_flow)))
                sputnik_a = self.sputnik.itemAtPosition(4, 3).widget().text()
                sputnik_b = self.sputnik.itemAtPosition(4, 4).widget().text()
                if sputnik_b == '':
                    sputnik_b = sputnik_a
                if all([flow, sputnik_a, sputnik_b]):
                    self._calculate_kms_by_radiobutton(flow, sputnik_a, sputnik_b)


    def calculate_kms(self, value) -> None:
        one_side_flow = self.sputnik.itemAtPosition(2, 1).widget().text()
        two_side_flow = self.sputnik.itemAtPosition(4, 1).widget().text()

        if all([self.radio_button1.isChecked(), one_side_flow]):
            sputnik_flow = one_side_flow
            sputnik_a = self.sputnik.itemAtPosition(2, 3).widget().text()
            sputnik_b = self.sputnik.itemAtPosition(2, 4).widget().text()
            if sputnik_b == '':
                sputnik_b = sputnik_a
            if all([one_side_flow, sputnik_a, sputnik_b]):
                self._calculate_kms_by_radiobutton(sputnik_flow, sputnik_a, sputnik_b)

        elif all([self.radio_button2.isChecked(), one_side_flow, two_side_flow]):
            flow = str(max(float(one_side_flow),  float(two_side_flow)))
            sputnik_a = self.sputnik.itemAtPosition(4, 3).widget().text()
            sputnik_b = self.sputnik.itemAtPosition(4, 4).widget().text()
            if sputnik_b == '':
                sputnik_b = sputnik_a
            if all([flow, sputnik_a, sputnik_b]):
                self._calculate_kms_by_radiobutton(flow, sputnik_a, sputnik_b)


    def _calculate_kms_by_radiobutton(self, flow, a, b) -> None:
        sputnik_a, sputnik_b = a, b
        if all([flow, sputnik_a, sputnik_b]):
            sputnik_flow = float(flow)
            sputnik_a, sputnik_b = int(sputnik_a), int(sputnik_b)
            rows = self.get_all_rows()
            for i in range(len(rows) - 1, 0, -1):
                pass_flow = rows[i-1].itemAtPosition(0, 3).widget().text()
                branch_flow = rows[i].itemAtPosition(0, 3).widget().text()
                main_a = rows[i].itemAtPosition(0, 10).widget().text()
                main_b = rows[i].itemAtPosition(0, 11).widget().text()
                if not main_b or main_b == '':
                    main_b = main_a

                if all([pass_flow, branch_flow, main_a, main_b]):
                    main_a, main_b = int(main_a), int(main_b)
                    pass_flow = float(pass_flow)
                    branch_flow = float(branch_flow)
                    Fk = (main_a / 1_000) * (main_b / 1_000)
                    Fs = (sputnik_a / 1_000) * (sputnik_b / 1_000)
                    pass_flow_relation = sputnik_flow / pass_flow
                    branch_flow_relation = sputnik_flow / branch_flow

                    try:
                        kms_1 = (1.55 * pass_flow_relation - pow(pass_flow_relation, 2))
                        pass_kms = kms_1 / (pow(1 - pass_flow_relation, 2) * pow(Fk / Fk, 2))
                        pass_kms = '{:.3f}'.format(round(pass_kms, 3))
                        rows[i].itemAtPosition(0, 8).widget().setText(pass_kms)
                    except ZeroDivisionError:
                        rows[i].itemAtPosition(0, 8).widget().setText('0')

                    try:
                        if (Fs / Fk <= 0.35) and (branch_flow_relation <= 1):
                            A = 1
                        elif branch_flow_relation <= 0.4:
                            A = 0.9 * (1 - branch_flow_relation)
                        else:
                            A = 0.55
                        kms_2 = A * (1 + pow(branch_flow_relation * (Fk / Fs), 2) - 2 * pow(1 - branch_flow_relation, 2))
                        branch_kms = kms_2 / pow(branch_flow_relation * (Fk / Fs), 2)

                        if i == len(rows) - 1:
                            rows[i].itemAtPosition(0, 9).widget().setText('3.7')
                        else:
                            branch_kms = '{:.3f}'.format(round(branch_kms, 3))
                            rows[i].itemAtPosition(0, 9).widget().setText(str(branch_kms))
                        self.last_row.itemAtPosition(0, 9).widget().setText('0')
                    except ZeroDivisionError:
                        rows[i].itemAtPosition(0, 9).widget().setText('')
                else:
                    rows[i].itemAtPosition(0, 8).widget().setText('')
                    rows[i].itemAtPosition(0, 9).widget().setText('')
        else:
            rows = self.get_main_rows()
            for i in range(len(rows)):
                rows[i].itemAtPosition(0, 8).widget().setText('')


    def calculate_pass_pressure(self, value) -> None:
        rows = self.get_all_rows()
        for row in rows:
            velocity = row.itemAtPosition(0, 12).widget().text()
            kms = row.itemAtPosition(0, 8).widget().text()
            temperature = self.temperature_widget.text()
            if all([kms, temperature, velocity]):
                velocity = float(velocity)
                kms = float(kms)
                temperature = float(temperature)
                result = kms * velocity * velocity * (353 / (273 + temperature)) / 2
                result = '{:.3f}'.format(round(result, 3))
                row.itemAtPosition(0, 18).widget().setText(result)
            else:
                row.itemAtPosition(0, 18).widget().setText('')


    def calculate_branch_pressure_by_radiobutton_1(self, checked) -> None:
        if checked:
            velocity = self.sputnik.itemAtPosition(2, 5).widget().text()
            if velocity:
                self._calculate_branch_pressure(velocity)


    def calculate_branch_pressure_by_radiobutton_2(self, checked) -> None:
        if checked:
            velocity_one = self.sputnik.itemAtPosition(2, 5).widget().text()
            velocity_two = self.sputnik.itemAtPosition(4, 5).widget().text()
            if velocity_one and not velocity_two:
                self._calculate_branch_pressure(velocity_one)
            elif all([velocity_one, velocity_two]):
                velocity = str(max(float(velocity_one),  float(velocity_two)))
                self._calculate_branch_pressure(velocity)
            else:
                rows = self.get_all_rows()
                for row in rows:
                    row.itemAtPosition(0, 19).widget().setText('')


    def calculate_branch_pressure(self, value) -> None:
        velocity_one = self.sputnik.itemAtPosition(2, 5).widget().text()
        velocity_two = self.sputnik.itemAtPosition(4, 5).widget().text()
        if velocity_one and not velocity_two:
            self._calculate_branch_pressure(velocity_one)
        elif all([velocity_one, velocity_two]):
            velocity = str(max(float(velocity_one),  float(velocity_two)))
            self._calculate_branch_pressure(velocity)


    def _calculate_branch_pressure(self, velocity) -> None:
        velocity = float(velocity)
        rows = self.get_all_rows()
        for row in rows:
            kms = row.itemAtPosition(0, 9).widget().text()
            temperature = self.temperature_widget.text()
            if all([kms, temperature]):
                kms = float(kms)
                temperature = float(temperature)
                result = kms * velocity * velocity * (353 / (273 + temperature)) / 2
                result = '{:.3f}'.format(round(result, 3))
                row.itemAtPosition(0, 19).widget().setText(result)
            else:
                row.itemAtPosition(0, 19).widget().setText('')


    def calculate_full_pressure_by_radiobutton_1(self, checked) -> None:
        if checked:
            klapan_full_pressure_1 = self.cell_3_13.text()
            if klapan_full_pressure_1:
                self._calculate_full_pressure(klapan_full_pressure_1)
            else:
                rows = self.get_all_rows()
                for row in rows:
                    row.itemAtPosition(0, 20).widget().setText('')


    def calculate_full_pressure_by_radiobutton_2(self, checked) -> None:
        if checked:
            klapan_full_pressure_1 = self.cell_3_13.text()
            klapan_full_pressure_2 = self.cell_5_13.text()
            if all([klapan_full_pressure_1, klapan_full_pressure_2]):
                klapan_full_pressure = str(max(float(klapan_full_pressure_1), float(klapan_full_pressure_2)))
                self._calculate_full_pressure(klapan_full_pressure)
            else:
                rows = self.get_all_rows()
                for row in rows:
                    row.itemAtPosition(0, 20).widget().setText('')


    def calculate_full_pressure(self, value) -> None:
        klapan_full_pressure_1 = self.cell_3_13.text()
        klapan_full_pressure_2 = self.cell_5_13.text()

        if all([self.radio_button1.isChecked(), klapan_full_pressure_1]):
            klapan_full_pressure = klapan_full_pressure_1
            self._calculate_full_pressure(klapan_full_pressure)
        elif all([self.radio_button2.isChecked(), klapan_full_pressure_1, klapan_full_pressure_2]):
            klapan_full_pressure = str(max(float(klapan_full_pressure_1), float(klapan_full_pressure_2)))
            self._calculate_full_pressure(klapan_full_pressure)
        else:
            rows = self.get_all_rows()
            for row in rows:
                row.itemAtPosition(0, 20).widget().setText('')


    def _calculate_full_pressure(self, klapan_full_pressure) -> None:
        rows = self.get_main_rows()
        for i in range(len(rows) - 1, -1, -1):
            branch_pressure = rows[i].itemAtPosition(0, 19).widget().text()
            if all([klapan_full_pressure, branch_pressure]):
                branch_pressure = float(branch_pressure)
                klapan_full_pressure = float(klapan_full_pressure)

                all_pass_pressure = []
                for line in range(i, -1, -1):
                    pass_pressure = rows[line].itemAtPosition(0, 18).widget().text()
                    if pass_pressure:
                        all_pass_pressure.append(pass_pressure)
                sum_pass_pressure = sum(map(float, all_pass_pressure))

                all_linear_pressure = []
                for line in range(i, -1, -1):
                    linear_pressure = rows[line].itemAtPosition(0, 16).widget().text()
                    if linear_pressure:
                        all_linear_pressure.append(linear_pressure)
                sum_linear_pressure = sum(map(float, all_linear_pressure))

                if self.cap_type.currentText() != CONSTANTS.CAP.TYPES[-1] and self.cap_pressure.text():
                    cap_pressure = float(self.cap_pressure.text())
                else:
                    cap_pressure = 0

                result = klapan_full_pressure + branch_pressure + sum_pass_pressure + sum_linear_pressure + cap_pressure
                result = '{:.3f}'.format(round(result, 3))
                rows[i].itemAtPosition(0, 20).widget().setText(result)
            else:
                rows[i].itemAtPosition(0, 20).widget().setText('')
                self.last_row.itemAtPosition(0, 20).widget().setText('')


    def calculate_full_pressure_last_row(self, value) -> None:
        row = self.last_row
        linear_pressure = row.itemAtPosition(0, 16).widget().text()
        pass_pressure = row.itemAtPosition(0, 18).widget().text()
        branch_pressure = row.itemAtPosition(0, 19).widget().text()
        if all([linear_pressure != '', pass_pressure != '', branch_pressure != '']):
            linear_pressure = float(linear_pressure)
            pass_pressure = float(pass_pressure)
            branch_pressure = float(branch_pressure)

            if self.cap_type.currentText() != CONSTANTS.CAP.TYPES[-1] and self.cap_pressure.text():
                cap_pressure = float(self.cap_pressure.text())
            else:
                cap_pressure = 0
            result = branch_pressure + pass_pressure + linear_pressure + cap_pressure
            result = '{:.3f}'.format(round(result, 3))
            row.itemAtPosition(0, 20).widget().setText(result)
        else:
            row.itemAtPosition(0, 20).widget().setText('')


    def activate_channel_cap(self, state) -> None:
        if state == 2:
            self.channel_cap_widget.hide()
        else:
            self.channel_cap_widget.setVisible(True)


    def change_channel_cap_visibility(self, value) -> None:
        current_value = self.cap_type.currentText()
        if current_value == CONSTANTS.CAP.TYPES[1]:
            for i in (8, 9, 10):
                self.channel_cap_grid.itemAtPosition(0, i).widget().setVisible(True)
            for i in range(2, 8):
                self.channel_cap_grid.itemAtPosition(0, i).widget().hide()
        elif current_value in CONSTANTS.CAP.TYPES[2:4]:
            for i in (2, 3, 4, 5, 6, 7, 8, 9, 10):
                self.channel_cap_grid.itemAtPosition(0, i).widget().setVisible(True)
        elif current_value == CONSTANTS.CAP.TYPES[-1]:
            for i in range(2, 11):
                self.channel_cap_grid.itemAtPosition(0, i).widget().hide()
        else:
            for i in range(2, 8):
                self.channel_cap_grid.itemAtPosition(0, i).widget().hide()


    def set_channel_cap_tooltip(self, value) -> None:
        if value == CONSTANTS.CAP.TYPES[2]:
            path = os.path.join(basedir, CONSTANTS.CAP.TYPES_IMG[0])
            self.cap_type.setToolTip(f'<img src="{path}" width="250">')
        elif value == CONSTANTS.CAP.TYPES[3]:
            path = os.path.join(basedir, CONSTANTS.CAP.TYPES_IMG[1])
            self.cap_type.setToolTip(f'<img src="{path}" width="250">')
        else:
            self.cap_type.setToolTip('')


    def set_channel_cap_relations(self, value) -> None:
        if value in (CONSTANTS.CAP.TYPES[2:4]):
            self.relations.clear()
            self.relations.addItems(CONSTANTS.CAP.RELATIONS.get(value))
            self.relations.insertSeparator(2)
            self.relations.model().item(0).setEnabled(False)
            self.relations.model().item(1).setEnabled(False)


    def calculate_channel_cap(self, value) -> None:
        current_value = self.cap_type.currentText()
        if current_value in CONSTANTS.CAP.TYPES[1:4]:
            if self.get_main_rows():
                if current_value == CONSTANTS.CAP.TYPES[1]:
                    w = self.get_main_rows()[0].itemAtPosition(0, 12).widget().text()
                    D = self.get_main_rows()[0].itemAtPosition(0, 13).widget().text()
                    temperature = self.temperature_widget.text()
                    if all([D, w, temperature]):
                        kms = 1
                        D, w, temperature = float(D), float(w), float(temperature)
                        result  = kms * (353 / (273.15 + temperature)) * pow(w, 2) / 2
                        result = '{:.3f}'.format(round(result, 3))
                        self.cap_pressure.setText(result)
                        self._calculate_channel_cap(result)
                    else:
                        self.cap_pressure.setText('')
                elif current_value in CONSTANTS.CAP.TYPES[2:4]:
                    D = self.get_main_rows()[0].itemAtPosition(0, 13).widget().text()
                    h = self.input_h.text()
                    if D and h:
                        D, h = float(D), float(h)
                        relation = h / D
                        relation = '{:.2f}'.format(round(relation, 2))
                        self.fact_relation.setText(relation)
                    else:
                        self.fact_relation.setText('')

                    w = self.get_main_rows()[0].itemAtPosition(0, 12).widget().text()
                    user_relation = self.relations.currentText()
                    kms_data = CONSTANTS.CAP.RELATIONS.get(current_value)
                    kms = kms_data.get(user_relation, False)
                    temperature = self.temperature_widget.text()
                    if all([D, w, kms, temperature]):
                        w, temperature = float(w), float(temperature)
                        result  = kms * (353 / (273.15 + temperature)) * pow(w, 2) / 2
                        result = '{:.3f}'.format(round(result, 3))
                        self.cap_pressure.setText(result)
                        self._calculate_channel_cap(result)
                    else:
                        self.cap_pressure.setText('')


    def _calculate_channel_cap(self, pressure) -> None:
        pressure = float(pressure)
        rows = self.get_all_rows()
        for row in rows:
            loss = row.itemAtPosition(0, 20).widget().text()
            if loss:
                loss = float(loss)
                result = loss + pressure
                result = '{:.3f}'.format(round(result, 3))
                row.itemAtPosition(0, 20).widget().setText(result)
            else:
                row.itemAtPosition(0, 20).widget().setText(str(pressure))


    def get_all_rows(self) -> list:
        sorted_rows = sorted([widget for widget in self.scroll_area.findChildren(QGridLayout)], key=lambda w: self.main_box.indexOf(w.parent()))
        rows = list(filter(lambda w: 'row' in w.objectName(), sorted_rows))
        return rows


    def get_main_rows(self) -> list:
        sorted_rows = sorted([widget for widget in self.scroll_area.findChildren(QGridLayout)], key=lambda w: self.main_box.indexOf(w.parent()))
        rows = list(filter(lambda w: w.objectName() == 'row', sorted_rows))
        return rows


    def get_sum_all_rows_int(self) -> int:
        return len(self.get_all_rows())


    def get_sum_all_rows_str(self) -> int:
        return str(len(self.get_all_rows()))


    def remove_all_main_rows(self) -> None:
        last_row = self.last_row
        parent_widget = last_row.parent()
        parent_widget.setParent(None)
        parent_widget.deleteLater()
        rows = self.get_main_rows()
        for row in rows:
            parent_widget = row.parent()
            parent_widget.setParent(None)
            parent_widget.deleteLater()
            self.rows_count -= 1


    def clean_all_input_data(self) -> None:
        init_data = self.init_data_layout
        for i in (0, 1, 2, 3, 5):
            init_data.itemAtPosition(i, 1).widget().setText('')
        self.klapan_widget.setCurrentIndex(-1)

        sputnik_data = self.sputnik
        self.klapan_flow.setText('')
        for i in (1, 2, 3, 4, 11):
            sputnik_data.itemAtPosition(2, i).widget().setText('')
        for i in (1, 2, 11):
            sputnik_data.itemAtPosition(4, i).widget().setText('')
        self.radio_button1.setChecked(True)

        self.tab_widget.setTabVisible(1, False)
        deflector = self.deflector
        for i in range(9):
            deflector.itemAtPosition(i, 1).widget().setText('')

        rows = self.get_all_rows()
        for row in rows:
            for i in range(22):
                row.itemAtPosition(0, i).widget().setText('')


    def open_manual(self):
        QMessageBox.information(self, 'Информация', 'Руководство находится в разработке')
        # if platform.system() == 'Windows':
        #     os.startfile(os.path.join(basedir, 'natural_air_system_manual.pdf'))


    def show_about(self) -> None:
        QMessageBox.information(self, 'О программе', CONSTANTS.ABOUT)


    def check_updates(self) -> None:
        current_version = tuple(map(int, version.split('.')))
        url = 'https://api.github.com/repos/polnikov/air-system/releases/latest'
        response = requests.get(url)
        data = response.json()
        try:
            latest_version = tuple(map(int, data['tag_name'].replace('v', '').split('.')))
            download_url = data['assets'][0]['browser_download_url']
            if latest_version > current_version:
                reply = QMessageBox.information(
                    self, 'Проверка обновления',
                    f'Новая версия {".".join(map(str, latest_version))} доступна для загрузки. Загрузить сейчас?',
                    QMessageBox.Yes | QMessageBox.No
                )
                if reply == QMessageBox.Yes:
                    self.download_file(download_url)
            else:
                QMessageBox.information(self, 'Проверка обновления', 'Вы используете последнюю версию')
        except KeyError:
            QMessageBox.information(self, 'Проверка обновления', 'Проверка обновлений временно недоступна. Попробуйте, пожалуйста, попозже.')


    def download_file(self, url):
        webbrowser.open(url)


    def _get_data_for_save(self) -> dict:
        data = {}

        init_data = []
        temperature = self.temperature_widget.text()
        surface = self.surface_widget.text()
        floor_height = self.floor_height_widget.text()
        channel_height = self.channel_height_widget.text()
        klapan_1 = self.klapan_input.text()
        klapan_2 = self.klapan_widget.currentText()

        init_data.append(temperature)
        init_data.append(surface)
        init_data.append(floor_height)
        init_data.append(channel_height)
        init_data.append(klapan_1)
        init_data.append(klapan_2)
        data['init_data'] = init_data

        sputnik_data = []
        klapan_flow = self.klapan_flow.text()
        one_side = [self.sputnik.itemAtPosition(2, i).widget().text() for i in (1, 2, 3, 4, 11)]
        sputnik_data.append(klapan_flow)
        sputnik_data.append({'one_side': one_side})

        two_side = [self.sputnik.itemAtPosition(4, i).widget().text() for i in (1, 2, 11)]
        sputnik_data.append({'two_side': two_side})

        if self.radio_button1.isChecked():
            sputnik_data.append({'is_checked': 1})
        else:
            sputnik_data.append({'is_checked': 2})
        data['sputnik_data'] = sputnik_data

        cap = self.cap_type.currentText()
        if cap == CONSTANTS.CAP.TYPES[-1]:
            wind = self.deflector.itemAtPosition(0, 1).widget().text()
            flow = self.deflector.itemAtPosition(2, 1).widget().text()
            data['deflector'] = [wind, flow]
        elif cap == CONSTANTS.CAP.TYPES[1]:
            data['cap_0'] = cap
        elif cap in CONSTANTS.CAP.TYPES[2:4]:
            h = self.channel_cap_grid.itemAtPosition(0, 2).widget().text()
            relation = self.relations.currentText()
            data['cap_1'] = [cap, h, relation]

        last_row = [self.last_row.itemAtPosition(0, i).widget().text() for i in (1, 2, 8, 10, 11)]
        data['last_row'] = last_row

        rows = []
        for row in self.get_main_rows():
            row_data = [row.itemAtPosition(0, i).widget().text() for i in (1, 2, 10, 11)]
            rows.append(row_data)
        data['rows'] = rows

        return data


    def _get_data_for_export(self) -> dict:
        data = {}

        init_data = []
        init_layout = self.init_data_layout
        for i in range(4):
            init_data.append(init_layout.itemAtPosition(i, 1).widget().text())

        if not self.klapan_input.text():
            init_data.append(f'{self.klapan_widget.currentText()} / {CONSTANTS.INIT_DATA.KLAPAN_ITEMS.get(self.klapan_widget.currentText())} м3/ч')
        else:
            init_data.append(f'- / {self.klapan_input.text()} м3/ч')

        init_data.append(self.klapan_flow.text())
        init_data.append(self.sputnik.itemAtPosition(1, 13).widget().text())
        data['init_data'] = init_data

        sputnik_data = {}
        sputnik_layout = self.sputnik
        sputnik_data['headers'] = [sputnik_layout.itemAtPosition(0, i).widget().text() for i in range(13)]
        if self.radio_button1.isChecked():
            sputnik_data['line_1'] = [sputnik_layout.itemAtPosition(2, i).widget().text() for i in range(13)]
            sputnik_data['line_2'] = sputnik_layout.itemAtPosition(3, 13).widget().text()
        else:
            sputnik_data['line_1'] = [sputnik_layout.itemAtPosition(2, i).widget().text() for i in range(13)]
            sputnik_data['line_2'] = sputnik_layout.itemAtPosition(3, 13).widget().text()
            sputnik_data['line_3'] = [sputnik_layout.itemAtPosition(4, i).widget().text() for i in range(13)]
            sputnik_data['line_4'] = sputnik_layout.itemAtPosition(5, 13).widget().text()
        data['sputnik_data'] = sputnik_data

        main_data = {}
        rows = self.get_all_rows()
        main_data['num_rows'] = self.get_sum_all_rows_int()+1
        cap = self.cap_type.currentText()
        if cap != CONSTANTS.CAP.TYPES[-1]:
            main_data['num_cols'] = 21
            main_data['headers'] = [self.header.itemAtPosition(0, i).widget().text() for i in range(22) if i != 6]
            for i in range(len(rows)):
                main_data[f'line_{i}'] = [rows[i].itemAtPosition(0, j).widget().text() for j in range(22) if j != 6]
            data['main_data'] = main_data

            if cap == CONSTANTS.CAP.TYPES[1]:
                data['cap_0'] = [cap, self.cap_pressure.text()]
            elif cap in CONSTANTS.CAP.TYPES[2:4]:
                h = self.channel_cap_grid.itemAtPosition(0, 3).widget().text()
                data['cap_1'] = [cap, h, self.relations.currentText(), self.cap_pressure.text()]
        else:
            main_data['num_cols'] = 22
            main_data['headers'] = [self.header.itemAtPosition(0, i).widget().text() for i in range(22)]
            for i in range(len(rows)):
                main_data[f'line_{i}'] = [rows[i].itemAtPosition(0, j).widget().text() for j in range(22)]
            data['main_data'] = main_data

            deflector_data = []
            deflector_data.append([self.deflector.itemAtPosition(i, 0).widget().text() for i in range(9)])
            deflector_data.append([self.deflector.itemAtPosition(i, 1).widget().text() for i in range(9)])
            data['deflector_data'] = deflector_data

        return data


    def export(self) -> None:
        data = self._get_data_for_export()
        if data:
            doc = docx.Document()

            doc_style = doc.styles.add_style('DocStyle', 1)
            doc_style.font.name = 'Times New Roman'
            doc_style.font.size = Pt(12)
            doc.styles['Normal'].base_style = doc_style

            title_style = doc.styles.add_style('TitleStyle', 1)
            title_style.font.name = 'Times New Roman'
            title_style.font.size = Pt(12)
            title_style.font.bold = True
            title = doc.add_paragraph(CONSTANTS.EXPORT.TITLE, style='TitleStyle')
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            table_title = doc.styles.add_style('TableTitleStyle', 1)
            table_title.font.name = 'Times New Roman'
            table_title.font.size = Pt(12)
            table_title.font.bold = True

            header_style = doc.styles.add_style('HeaderStyle', 1)
            header_style.font.bold = True

            # setup fields
            sections = doc.sections
            for section in sections:
                section.orientation = WD_ORIENT.LANDSCAPE
                new_width, new_height = section.page_height, section.page_width
                section.page_width = new_width
                section.page_height = new_height
                section.left_margin = Cm(2)
                section.right_margin = Cm(1)
                section.top_margin = Cm(1)
                section.bottom_margin = Cm(1)
                section.page_width = Mm(420)
                section.page_height = Mm(297)

            # add table 1
            doc.add_paragraph(CONSTANTS.EXPORT.INIT_DATA.TITLE, style='TableTitleStyle')
            table_1 = doc.add_table(rows=7, cols=3)
            table_1.style = 'Table Grid'

            # setup table header, units and column widths
            header = table_1.columns[0].cells
            for row in range(7):
                header[row].text = CONSTANTS.EXPORT.INIT_DATA.HEADER[row]
            units = table_1.columns[2].cells
            for row in range(7):
                units[row].text = CONSTANTS.EXPORT.INIT_DATA.UNITS[row]
                units[row].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                units[row].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            widths = [110, 90, 27]
            for col in range(3):
                set_column_width(table_1.columns[col], widths[col])
            for row in range(7):
                set_row_height(table_1.rows[row], 7)

            init_data = data['init_data']
            values = table_1.columns[1].cells
            for row in range(7):
                values[row].text = init_data[row]
                values[row].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                values[row].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            doc.add_paragraph()

            # add table 2
            doc.add_paragraph(CONSTANTS.EXPORT.SPUTNIK.TITLE, style='TableTitleStyle')

            sputnik_data = data['sputnik_data']
            table_2 = doc.add_table(rows=len(sputnik_data), cols=13)
            table_2.style = 'Table Grid'

            header = table_2.rows[0].cells
            for col in range(13):
                header[col].text = sputnik_data['headers'][col]
                header[col].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                header[col].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            for row in range(1, len(sputnik_data)):
                set_row_height(table_2.rows[row], 7)

            for row in range(1, len(sputnik_data)):
                values = table_2.rows[row].cells
                if row in (1, 3):
                    for col in range(13):
                        line = f'line_{row}'
                        values[col].text = sputnik_data[line][col]
                        values[col].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        values[col].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                elif row in (2, 4):
                    line = f'line_{row}'
                    values[12].text = sputnik_data[line]
                    values[12].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    values[12].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            for row in range(1, len(sputnik_data)):
                set_row_height(table_2.rows[row], 7)

            if len(sputnik_data) == 3:
                a = table_2.cell(2, 0)
                b = table_2.cell(2, 11)
                a.merge(b)
            else:
                for i in (2, 4):
                    a = table_2.cell(i, 0)
                    b = table_2.cell(i, 11)
                    a.merge(b)

            for col in range(13):
                set_column_width(table_2.columns[col], 25)
            doc.add_paragraph()

            # add table 3
            doc.add_paragraph(CONSTANTS.EXPORT.MAIN_TABLE.TITLE, style='TableTitleStyle')

            main_data = data['main_data']
            table_3 = doc.add_table(rows=main_data['num_rows'], cols=main_data['num_cols'])
            table_3.style = 'Table Grid'

            header = table_3.rows[0].cells
            for col in range(main_data['num_cols']):
                header[col].text = main_data['headers'][col]
                header[col].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                header[col].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                header[col].paragraphs[0].runs[0].font.size = Pt(10)

            for row in range(1, main_data['num_rows']):
                set_row_height(table_3.rows[row], 7)

            for row in range(1, main_data['num_rows']):
                values = table_3.rows[row].cells
                for col in range(main_data['num_cols']):
                    line = f'line_{row - 1}'
                    values[col].text = main_data[line][col]
                    values[col].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    values[col].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    values[col].paragraphs[0].runs[0].font.size = Pt(10)

            if  main_data['num_cols'] == 22:
                a = table_3.cell(2, 6)
                b = table_3.cell(main_data['num_rows'] - 1, 6)
                a.merge(b)

            set_column_width(table_3.columns[0], 15)
            set_column_width(table_3.columns[1], 15)
            set_column_width(table_3.columns[main_data['num_cols'] - 1], 25)
            doc.add_paragraph()

            # add formula
            doc.add_paragraph(CONSTANTS.EXPORT.FORMULA, style='TableTitleStyle')
            doc.add_paragraph().add_run().add_picture(os.path.join(basedir, CONSTANTS.EXPORT.FORMULA_IMG), width=Cm(12))
            doc.add_paragraph()

            if data.get('cap_0'):
                # add table 4
                doc.add_paragraph(CONSTANTS.EXPORT.CAP.TITLE, style='TableTitleStyle')

                cap_data = data['cap_0']
                table_4 = doc.add_table(rows=2, cols=2)
                table_4.style = 'Table Grid'

                widths = [50, 20]
                for col in range(2):
                    set_column_width(table_4.columns[col], widths[col])
                for row in range(2):
                    set_row_height(table_4.rows[row], 7)

                header = table_4.columns[0].cells
                for row in range(2):
                    header[row].text = CONSTANTS.EXPORT.CAP.HEADER_1[row]

                values = table_4.columns[1].cells
                for row in range(2):
                    values[row].text = cap_data[row]
                    values[row].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    values[row].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                doc.add_paragraph()

            if data.get('cap_1'):
                # add table 4
                doc.add_paragraph(CONSTANTS.EXPORT.CAP.TITLE, style='TableTitleStyle')

                cap_data = data['cap_1']
                table_4 = doc.add_table(rows=5, cols=4)
                table_4.style = 'Table Grid'
                table_4.allow_autofit = False
                table_4.alignment = WD_TABLE_ALIGNMENT.LEFT

                widths = [90, 37, 20, 25]
                for col in range(4):
                    set_column_width(table_4.columns[col], widths[col])
                for row in range(5):
                    set_row_height(table_4.rows[row], 7)

                header = table_4.columns[0].cells
                for row in range(5):
                    header[row].text = CONSTANTS.EXPORT.CAP.HEADER_2[row]

                units = table_4.columns[2].cells
                for row in range(5):
                    units[row].text = CONSTANTS.EXPORT.CAP.UNITS_2[row]
                    units[row].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    units[row].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                values = table_4.columns[1].cells
                for row in range(5):
                    match row:
                        case 2:
                            values[row].text = cap_data[2].split(':')[0]
                        case 3:
                            values[row].text = cap_data[2].split(':')[1]
                        case 4:
                            values[row].text = cap_data[3]
                        case _:
                            values[row].text = cap_data[row]
                    values[row].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    values[row].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                a = table_4.cell(0, 3)
                b = table_4.cell(4, 3)
                a.merge(b)
                delete_paragraph(table_4.cell(0, 3).paragraphs[0])

                if data['cap_1'][0] == 'Зонт':
                    table_4.cell(0, 3).add_paragraph().add_run().add_picture(os.path.join(basedir, CONSTANTS.CAP.TYPES_IMG[0]), height=Cm(3.5))
                else:
                    table_4.cell(0, 3).add_paragraph().add_run().add_picture(os.path.join(basedir, CONSTANTS.CAP.TYPES_IMG[1]), height=Cm(3.5))

            if data.get('deflector_data'):
                # add table 5
                doc.add_paragraph(CONSTANTS.EXPORT.CAP.DEFLECTOR_TITLE, style='TableTitleStyle')

                deflector_data = data['deflector_data']
                table_5 = doc.add_table(rows=9, cols=2)
                table_5.style = 'Table Grid'

                header = table_5.columns[0].cells
                for row in range(9):
                    header[row].text = deflector_data[0][row]

                widths = [120, 27]
                for col in range(2):
                    set_column_width(table_5.columns[col], widths[col])
                for row in range(9):
                    set_row_height(table_5.rows[row], 7)

                values = table_5.columns[1].cells
                for row in range(9):
                    values[row].text = deflector_data[1][row]
                    values[row].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    values[row].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            now = datetime.now()
            date_time = now.strftime('%d_%m_%y_%H_%M')
            if self.current_file_path:
                file_name = f'{self.current_file_path}_{date_time}'
            else:
                file_name = f'Без названия_{date_time}'
            file_name = file_name.replace('.json', '')
            doc.save(f'{file_name}.docx')
            QMessageBox.information(self, 'Информация', 'Расчёт успешно экспортирован')
            os.startfile(os.path.join(basedir, f'{file_name}.docx'))
        else:
            QMessageBox.critical(self, 'Ошибка', 'Пока нечего экспортировать')


    def save_as(self) -> None:
        data = self._get_data_for_save()
        save_dir = QStandardPaths.writableLocation(QStandardPaths.StandardLocation.DocumentsLocation)
        file_name, _ = QFileDialog.getSaveFileName(self, 'Сохранить расчёт', save_dir, 'JSON (*.json)')
        if file_name:
            self.current_file_path = file_name
            self.setWindowTitle(f'{self.app_title} | {file_name}')

            with open(file_name, 'w') as file:
                json.dump(data, file)


    def save(self) -> None:
        if self.current_file_path is None:
            self.save_as()
        else:
            data = self._get_data_for_save()
            with open(self.current_file_path, 'w') as file:
                json.dump(data, file)


    def auto_save(self) -> None:
        if self.current_file_path:
            self.save()
        else:
            self.save_as()


    def open(self) -> None:
        options = QFileDialog.Options()
        open_dir = QStandardPaths.writableLocation(QStandardPaths.StandardLocation.DocumentsLocation)
        file_name, _ = QFileDialog.getOpenFileName(self, 'Выберите файл расчёта', open_dir, 'JSON файл (*.json);;Все файлы (*)', options=options)
        self._open_file(file_name)


    def _open_file(self, file_name) -> None:
        if file_name:
            try:
                with open(file_name) as f:
                    data = json.load(f)

                    progress = QProgressDialog('Импорт данных', None, 0, 100, self)
                    progress.setWindowTitle('Заполнение данных...')
                    progress.setWindowModality(Qt.WindowModality.WindowModal)
                    progress.setMinimumDuration(0)
                    progress.setValue(0)
                    progress.setMaximum(100)
                    progress.show()

                    # prepare main table
                    num_rows = len(data['rows'])
                    self.remove_all_main_rows()
                    self.clean_all_input_data()

                    self.main_box.addWidget(self.create_last_row())
                    for i in range(num_rows):
                        self.main_box.insertWidget(2, self.create_row())
                    self.last_row.itemAtPosition(0, 0).widget().setText(self.get_sum_all_rows_str())
                    self.change_dimensions_cells_in_table()

                    progress.setValue(progress.value() + 20)

                    # init data
                    init_data = data['init_data']
                    self.temperature_widget.setText(init_data[0])
                    self.surface_widget.setText(init_data[1])
                    self.floor_height_widget.setText(init_data[2])
                    self.channel_height_widget.setText(init_data[3])
                    self.klapan_input.setText(init_data[4])
                    self.klapan_widget.setCurrentText(init_data[5])

                    progress.setValue(progress.value() + 10)

                    # sputnik data
                    self.klapan_flow.setText(data['sputnik_data'][0])
                    one_side = data['sputnik_data'][1]['one_side']
                    for i, j in enumerate((1, 2, 3, 4, 11)):
                        self.sputnik.itemAtPosition(2, j).widget().setText(one_side[i])

                    progress.setValue(progress.value() + 10)

                    two_side = data['sputnik_data'][2]['two_side']
                    for i, j in enumerate((1, 2, 11)):
                        self.sputnik.itemAtPosition(4, j).widget().setText(two_side[i])

                    progress.setValue(progress.value() + 10)

                    if data['sputnik_data'][3]['is_checked'] == 1:
                        self.radio_button1.setChecked(True)
                    else:
                        self.radio_button2.setChecked(True)

                    progress.setValue(progress.value() + 10)

                    # last row data
                    last_row = data['last_row']
                    for i, j in enumerate((1, 2, 8, 10, 11)):
                        self.last_row.itemAtPosition(0, j).widget().setText(last_row[i])

                    progress.setValue(progress.value() + 10)

                    # table data
                    rows_data = data['rows']
                    rows = self.get_main_rows()
                    for i in range(num_rows):
                        for j, k in enumerate((1, 2, 10, 11)):
                            rows[i].itemAtPosition(0, k).widget().setText(rows_data[i][j])

                    progress.setValue(progress.value() + 20)

                    # deflector data
                    if data.get('deflector', False):
                        self.deflector.itemAtPosition(0, 1).widget().setText(data['deflector'][0])
                        self.deflector.itemAtPosition(2, 1).widget().setText(data['deflector'][1])
                        self.tab_widget.setTabVisible(1, True)
                        self.cap_type.setCurrentIndex(4)
                    # сap data
                    else:
                        self.tab_widget.setTabVisible(1, False)
                        if data.get('cap_0', False):
                            self.cap_type.setCurrentText(CONSTANTS.CAP.TYPES[1])
                        else:
                            self.cap_type.setCurrentText(data['cap_1'][0])
                            self.channel_cap_grid.itemAtPosition(0, 2).widget().setText(data['cap_1'][1])
                            self.relations.setCurrentText(data['cap_1'][2])

                    progress.setValue(progress.value() + 10)

                self.current_file_path = file_name
                self.setWindowTitle(f'{self.app_title} | {file_name}')
                progress.close()
                self.add_recent_file(file_name)
            except FileNotFoundError:
                QMessageBox.critical(self, 'Ошибка', 'Такого файла не существует или он был перемещен')
                self.recent_files.remove(file_name)
                self.settings.setValue('recentFiles', list(self.recent_files))
                self.update_recent_files_menu()
            except Exception as e:
                QMessageBox.critical(self, 'Ошибка', f'Не удалось открыть файл:\n{e}')
        else:
            QMessageBox.critical(self, 'Ошибка', 'Что то пошло не так...')


    def open_recent_file(self, file_path) -> None:
        self._open_file(file_path)


    def add_recent_file(self, file_path) -> None:
        if file_path in self.recent_files:
            self.recent_files.remove(file_path)
        self.recent_files.append(file_path)
        self.settings.setValue('recentFiles', list(self.recent_files))
        self.update_recent_files_menu()


    def load_recent_files(self) -> None:
        recent_files = self.settings.value('recentFiles', [])
        if recent_files:
            self.recent_files.extend(recent_files)


    def update_recent_files_menu(self) -> None:
        self.recent_files_menu.clear()
        recent_files = self.settings.value('recentFiles', [])
        for i, file_path in enumerate(recent_files):
            action = QAction(file_path, self)
            action.setIcon(QIcon(os.path.join(basedir, CONSTANTS.RECENT_FILES_ICONS[i])))
            action.triggered.connect(partial(self.open_recent_file, file_path))
            self.recent_files_menu.addAction(action)

        self.recent_files_menu.addSeparator()

        clear_action = QAction('Очистить список', self)
        clear_action.setIcon(QIcon(os.path.join(basedir, CONSTANTS.RECENT_FILES_ICONS[-2])))
        clear_action.triggered.connect(self.clear_recent_files)
        self.recent_files_menu.addAction(clear_action)


    def clear_recent_files(self) -> None:
        self.recent_files.clear()
        self.settings.remove('recentFiles')
        self.update_recent_files_menu()


    def closeEvent(self, event) -> None:
        self.settings.setValue('recentFiles', list(self.recent_files))

        if self.current_file_path is None:
            reply = QMessageBox.question(
                self,
                'Подтверждение',
                '''<html><center>Вы уверены, что хотите закрыть программу?<br><font color="red">Несохраненный расчет будет потерян</font></center></html>
                ''',
                QMessageBox.Yes | QMessageBox.No, QMessageBox.No
            )

            if reply == QMessageBox.No:
                self.save()
                event.accept()
                event.ignore()
            else:
                event.accept()
        else:
            reply = QMessageBox.question(self, 'Подтверждение', 'Сохранить текущие изменения?', QMessageBox.Yes | QMessageBox.No, QMessageBox.No)

            if reply == QMessageBox.No:
                event.accept()
            else:
                self.save
                event.accept()
        super().closeEvent(event)


def set_column_width(column, width) -> None:
    column.width = Mm(width)
    for cell in column.cells:
        cell.width = Mm(width)


def set_row_height(row, height) -> None:
    row.height = Mm(height)
    for cell in row.cells:
        cell.height = Mm(height)


def set_row_text_bold(row) -> None:
    for cell in row.cells:
        cell.bold = True


def delete_paragraph(paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


if __name__ == '__main__':
    import sys
    app = QApplication(sys.argv)
    app.setFont(QFont('Consolas', 10))
    app.setStyleSheet('QMessageBox { messagebox-text-interaction-flags: 5; font-size: 13px; }')
    app.setStyle('windowsvista')
    window = MainWindow()
    window.setWindowIcon(QIcon(os.path.join(basedir, 'app.ico')))
    window.setIconSize(QSize(15, 15))
    window.setStyleSheet('QMainWindow::title { background: black; }')
    window.show()
    sys.exit(app.exec())
